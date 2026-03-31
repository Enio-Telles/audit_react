from __future__ import annotations

import json
from io import BytesIO
from pathlib import Path

import polars as pl
from fastapi.testclient import TestClient

import api
import relatorio_fiscal


client = TestClient(api.app)


def _escrever_parquet(caminho: Path, dados: list[dict], schema: dict[str, pl.DataType] | None = None) -> None:
    caminho.parent.mkdir(parents=True, exist_ok=True)
    if dados:
        pl.DataFrame(dados).write_parquet(caminho, compression="zstd")
    else:
        pl.DataFrame(schema=schema or {}).write_parquet(caminho, compression="zstd")


def _preparar_cnpj_teste(base_dir: Path, cnpj: str = "12345678000190") -> Path:
    diretorio_cnpj = base_dir / cnpj
    for pasta in ["extraidos", "silver", "parquets", "edicoes", "exportacoes"]:
        (diretorio_cnpj / pasta).mkdir(parents=True, exist_ok=True)

    # Tabela base para leitura/exportacao
    _escrever_parquet(
        diretorio_cnpj / "parquets" / "produtos.parquet",
        [
            {
                "id_produto": 1,
                "descricao": "Produto A",
                "ncm": "22030000",
                "cest": None,
                "unidade_principal": "UN",
                "qtd_total_nfe": 3,
                "valor_total": 150.0,
                "tipo": "ambos",
            },
            {
                "id_produto": 2,
                "descricao": "Produto B",
                "ncm": "22030000",
                "cest": None,
                "unidade_principal": "UN",
                "qtd_total_nfe": 2,
                "valor_total": 90.0,
                "tipo": "compra",
            },
        ],
    )

    # Parquets minimos para dependencias do pipeline
    _escrever_parquet(
        diretorio_cnpj / "parquets" / "produtos_unidades.parquet",
        [],
        schema={
            "id_produto": pl.Int64,
            "descricao": pl.String,
            "ncm": pl.String,
            "cest": pl.String,
            "gtin": pl.String,
            "unid_compra": pl.String,
            "unid_venda": pl.String,
            "qtd_nfe_compra": pl.Int64,
            "qtd_nfe_venda": pl.Int64,
            "qtd_efd": pl.Int64,
            "valor_total_compra": pl.Float64,
            "valor_total_venda": pl.Float64,
        },
    )

    _escrever_parquet(
        diretorio_cnpj / "parquets" / "produtos_agrupados.parquet",
        [
            {
                "id_agrupado": "id_agrupado_1",
                "descricao_padrao": "Produto A",
                "ncm_padrao": "22030000",
                "cest_padrao": None,
                "ids_membros": "[1]",
                "qtd_membros": 1,
                "qtd_total_nfe": 3,
                "valor_total": 150.0,
                "unid_compra": "UN",
                "unid_venda": "UN",
                "origem": "manual",
                "criado_em": "2026-03-30T00:00:00+00:00",
                "editado_em": "2026-03-30T00:00:00+00:00",
                "status": "ativo",
            }
        ],
    )

    _escrever_parquet(
        diretorio_cnpj / "parquets" / "fatores_conversao.parquet",
        [
            {
                "id_agrupado": "id_agrupado_1",
                "descricao_padrao": "Produto A",
                "unid_compra": "UN",
                "unid_venda": "UN",
                "unid_ref": "UN",
                "fator_compra_ref": 1.0,
                "fator_venda_ref": 1.0,
                "origem_fator": "calculado",
                "status": "ok",
                "editado_em": None,
            }
        ],
    )

    _escrever_parquet(
        diretorio_cnpj / "parquets" / "produtos_final.parquet",
        [
            {
                "id_agrupado": "id_agrupado_1",
                "descricao_padrao": "Produto A",
                "ncm_padrao": "22030000",
                "cest_padrao": None,
                "unid_ref": "UN",
                "fator_compra_ref": 1.0,
                "fator_venda_ref": 1.0,
                "qtd_total_nfe": 3,
                "valor_total": 150.0,
                "ids_membros": "[1]",
                "qtd_membros": 1,
                "status_conversao": "ok",
                "status_agregacao": "ativo",
            }
        ],
    )

    _escrever_parquet(
        diretorio_cnpj / "parquets" / "id_agrupados.parquet",
        [
            {
                "id_produto": 1,
                "id_agrupado": "id_agrupado_1",
                "descricao_original": "Produto A",
                "descricao_padrao": "Produto A",
            }
        ],
    )

    _escrever_parquet(
        diretorio_cnpj / "parquets" / "nfe_entrada.parquet",
        [],
        schema={
            "chave_nfe": pl.String,
            "id_agrupado": pl.String,
            "data_emissao": pl.String,
            "cfop": pl.String,
            "quantidade": pl.Float64,
            "unidade": pl.String,
            "qtd_ref": pl.Float64,
            "valor_unitario": pl.Float64,
            "valor_total": pl.Float64,
            "cnpj_emitente": pl.String,
        },
    )

    _escrever_parquet(
        diretorio_cnpj / "silver" / "item_unidades.parquet",
        [
            {
                "descricao": "Produto A",
                "unidade": "UN",
                "compras": 150.0,
                "qtd_compras": 30.0,
                "vendas": 90.0,
                "qtd_vendas": 18.0,
            }
        ],
    )

    _escrever_parquet(
        diretorio_cnpj / "silver" / "tb_documentos.parquet",
        [
            {
                "chave_documento": "DOC001",
                "tipo_documento": "nfe_entrada",
                "data_documento": "2026-03-01",
            }
        ],
    )

    _escrever_parquet(
        diretorio_cnpj / "extraidos" / "reg0000.parquet",
        [
            {
                "nome": "EMPRESA TESTE LTDA",
                "ie": "123456789",
                "cpf": "12345678901",
            }
        ],
    )

    return diretorio_cnpj


def _preparar_relatorio_cnpj(
    base_dir: Path,
    cnpj: str,
    *,
    contribuinte: str = "EMPRESA TESTE LTDA",
    dsf: str = "20263710400226",
    pdf_det_invalido: bool = False,
    manifestacao: str = "Nao apresentou manifestacao",
    manifestacoes: dict | None = None,
    arquivos_notificacao_incluidos: list[str] | None = None,
) -> Path:
    diretorio_relatorio = base_dir / cnpj / "relatorio"
    diretorio_relatorio.mkdir(parents=True, exist_ok=True)

    payload = {
        "cnpj": cnpj,
        "contribuinte": contribuinte,
        "ie": "123456789",
        "dsf": dsf,
        "notificacao_det": "DET-001",
        "manifestacao": manifestacao,
        "contatos_realizados": "Contato realizado por e-mail.",
        "decisao_fiscal": "Encaminhar para acao fiscal.",
        "desfecho": "Pendencia mantida.",
    }

    if manifestacoes is not None:
        payload["manifestacoes"] = manifestacoes
    if arquivos_notificacao_incluidos is not None:
        payload["arquivos_notificacao_incluidos"] = arquivos_notificacao_incluidos

    if pdf_det_invalido:
        payload["pdf_det"] = "/home/ubuntu/upload/det_inexistente.pdf"

    (diretorio_relatorio / "dados.json").write_text(
        json.dumps(payload, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

    return diretorio_relatorio


def _criar_pdf_fake(caminho: Path) -> None:
    caminho.parent.mkdir(parents=True, exist_ok=True)
    caminho.write_bytes(b"%PDF-1.4\n%fake\n")


def _escapar_texto_pdf(texto: str) -> str:
    """Escapa texto ASCII simples para o content stream do PDF."""
    texto_normalizado = (
        texto.replace("\u2014", "-")
        .replace("\u2013", "-")
        .replace("\u2018", "'")
        .replace("\u2019", "'")
        .replace("\u201c", '"')
        .replace("\u201d", '"')
        .replace("\u00a0", " ")
        .encode("latin-1", errors="replace")
        .decode("latin-1")
    )
    return texto_normalizado.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def _criar_pdf_textual(caminho: Path, texto: str) -> None:
    """Cria um PDF textual minimo sem depender de reportlab."""
    caminho.parent.mkdir(parents=True, exist_ok=True)
    linhas = [linha for linha in texto.splitlines() if linha.strip()] or [""]
    comandos_texto = ["BT", "/F1 12 Tf", "72 720 Td"]
    for indice, linha in enumerate(linhas):
        texto_pdf = _escapar_texto_pdf(linha)
        if indice > 0:
            comandos_texto.append("0 -14 Td")
        comandos_texto.append(f"({texto_pdf}) Tj")
    comandos_texto.append("ET")

    objetos = [
        "<< /Type /Catalog /Pages 2 0 R >>",
        "<< /Type /Pages /Count 1 /Kids [3 0 R] >>",
        "<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>",
        "<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    stream = "\n".join(comandos_texto) + "\n"
    objetos.append(f"<< /Length {len(stream.encode('latin-1'))} >>\nstream\n{stream}endstream")

    conteudo = bytearray(b"%PDF-1.4\n")
    offsets = [0]

    for indice, objeto in enumerate(objetos, start=1):
        offsets.append(len(conteudo))
        conteudo.extend(f"{indice} 0 obj\n{objeto}\nendobj\n".encode("latin-1"))

    inicio_xref = len(conteudo)
    conteudo.extend(f"xref\n0 {len(objetos) + 1}\n".encode("latin-1"))
    conteudo.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        conteudo.extend(f"{offset:010d} 00000 n \n".encode("latin-1"))

    trailer = (
        f"trailer\n<< /Size {len(objetos) + 1} /Root 1 0 R >>\n"
        f"startxref\n{inicio_xref}\n%%EOF"
    )
    conteudo.extend(trailer.encode("latin-1"))
    caminho.write_bytes(bytes(conteudo))


def _instalar_renderizador_pdf_teste(monkeypatch) -> None:
    """Substitui apenas o renderer externo por um PDF textual deterministico."""

    def _renderizar_pdf_teste(html_content: str, output_path: str) -> None:
        linhas = relatorio_fiscal._extrair_texto_simples_html(html_content)
        texto = "\n".join(linha for linha in linhas if linha.strip())
        _criar_pdf_textual(Path(output_path), texto)

    monkeypatch.setattr(relatorio_fiscal, "_renderizar_pdf_html_ou_fallback", _renderizar_pdf_teste)


def _criar_modelos_docx_teste(tmp_path: Path) -> tuple[Path, Path]:
    from docx import Document

    caminho_individual = tmp_path / "modelo_individual.docx"
    caminho_geral = tmp_path / "modelo_geral.docx"

    doc_individual = Document()
    doc_individual.add_paragraph("Governo do Estado de Rondônia\nSecretaria de Estado de Finanças\nCoordenadoria da Receita Estadual")
    doc_individual.add_paragraph("______________________________________________________________________")
    doc_individual.add_paragraph("RELATÓRIO FISCAL CONCLUSIVO")
    doc_individual.add_paragraph("1. Identificação da Ação Fiscal")
    doc_individual.add_paragraph("Trata-se de ação fiscal preliminar decorrente de Designação de Serviço Fiscal (DSF).")
    doc_individual.add_paragraph("2. Procedimentos Realizados")
    doc_individual.add_paragraph("Foram adotadas as seguintes providências:")
    doc_individual.add_paragraph("Consulta ao Sismonitora para identificação das pendências e status;")
    doc_individual.add_paragraph("Verificação cadastral e fiscal do contribuinte no sistema Visão 360º;")
    doc_individual.add_paragraph("Conferência do status das inconsistências no sistema Fisconforme;")
    doc_individual.add_paragraph("Verificação da existência de monitoramentos fiscais relacionados;")
    doc_individual.add_paragraph("Emissão de notificação formal via DET, fixando prazo para regularização.")
    doc_individual.add_paragraph("3. Manifestação do Contribuinte")
    doc_individual.add_paragraph("Após a notificação, verificou-se que o contribuinte:")
    doc_individual.add_paragraph("( ) Regularizou integralmente as pendências")
    doc_individual.add_paragraph("( ) Apresentou contestação")
    doc_individual.add_paragraph("( ) Solicitou prorrogação de prazo")
    doc_individual.add_paragraph("( ) Não apresentou manifestação")
    doc_individual.add_paragraph("[DESCREVER_AQUI_OS_CONTATOS_REALIZADOS]")
    doc_individual.add_paragraph("4. Análise Fiscal")
    doc_individual.add_paragraph("Foram examinadas as justificativas apresentadas.")
    doc_individual.add_paragraph("[DESCREVER_AQUI_A_DECISAO_FISCAL]")
    doc_individual.add_paragraph("5. Situação Final das Pendências")
    doc_individual.add_paragraph("Após a conclusão da ação fiscal, as pendências apresentaram os seguintes desfechos:")
    doc_individual.add_paragraph("[DESCREVER_AQUI_O_DESFECHO_FINAL]")
    doc_individual.add_paragraph("6. Conclusão")
    doc_individual.add_paragraph("Diante do exposto, considera-se concluída a presente ação fiscal preliminar.")
    doc_individual.add_paragraph("[LOCAL_E_DATA]")
    doc_individual.add_paragraph("[ESPAÇO RESERVADO PARA ASSINATURA DIGITAL]")
    doc_individual.add_paragraph("____________________________________________________")
    doc_individual.add_paragraph("[NOME_DO_AUDITOR]")
    doc_individual.add_paragraph("[CARGO_DO_AUDITOR]")
    doc_individual.add_paragraph("Matrícula: [MATRICULA]")
    tabela = doc_individual.add_table(rows=5, cols=2)
    linhas = [
        ("Contribuinte:", "[NOME_DO_CONTRIBUINTE]"),
        ("CNPJ:", "[CNPJ]"),
        ("Inscrição Estadual:", "[INSCRICAO_ESTADUAL]"),
        ("DSF:", "[NUMERO_DSF]"),
        ("Notificação DET nº:", "[NUMERO_DET]"),
    ]
    for indice, (rotulo, valor) in enumerate(linhas):
        tabela.rows[indice].cells[0].text = rotulo
        tabela.rows[indice].cells[1].text = valor
    doc_individual.save(caminho_individual)

    doc_geral = Document()
    doc_geral.add_paragraph("Governo do Estado de Rondônia\nSecretaria de Estado de Finanças\nCoordenadoria da Receita Estadual")
    doc_geral.add_paragraph("______________________________________________________________________")
    doc_geral.add_paragraph("RELATÓRIO GERAL CONSOLIDADO")
    doc_geral.add_paragraph("Fisconforme Não Cumprido — DSF nº [NUMERO_DSF]")
    doc_geral.add_paragraph("Total de contribuintes analisados: [TOTAL_CONTRIBUINTES]")
    doc_geral.add_paragraph("1. Objeto")
    doc_geral.add_paragraph("O presente relatório consolida os resultados da ação fiscal preliminar realizada no âmbito da DSF nº [NUMERO_DSF].")
    doc_geral.add_paragraph("2. Quadro Resumo dos Contribuintes")
    doc_geral.add_paragraph("[INSERIR_TABELA_RESUMO_AQUI]")
    doc_geral.add_paragraph("3. Procedimentos Comuns")
    doc_geral.add_paragraph("Para todos os contribuintes listados, foram adotadas as seguintes providências:")
    doc_geral.add_paragraph("Consulta ao Sismonitora para identificação das pendências e status;")
    doc_geral.add_paragraph("Verificação cadastral e fiscal do contribuinte no sistema Visão 360º;")
    doc_geral.add_paragraph("Conferência do status das inconsistências no sistema Fisconforme;")
    doc_geral.add_paragraph("Verificação da existência de monitoramentos fiscais relacionados;")
    doc_geral.add_paragraph("Emissão de notificação formal via DET, fixando prazo para regularização;")
    doc_geral.add_paragraph("Tentativa de contato por e-mail, WhatsApp e telefone.")
    doc_geral.add_paragraph("4. Análise Individual por Contribuinte")
    doc_geral.add_paragraph("[REPETIR_O_BLOCO_ABAIXO_PARA_CADA_CONTRIBUINTE]")
    doc_geral.add_paragraph("--------------------------------------------------")
    doc_geral.add_paragraph("[NUMERO_ORDEM]. [NOME_DO_CONTRIBUINTE]")
    doc_geral.add_paragraph("CNPJ: [CNPJ]")
    doc_geral.add_paragraph("Inscrição Estadual: [INSCRICAO_ESTADUAL]")
    doc_geral.add_paragraph("DSF: [NUMERO_DSF]")
    doc_geral.add_paragraph("Notificação DET nº: [NUMERO_DET]")
    doc_geral.add_paragraph("Contatos e Manifestação:")
    doc_geral.add_paragraph("[DESCREVER_AQUI_OS_CONTATOS_E_MANIFESTACAO]")
    doc_geral.add_paragraph("Decisão Fiscal:")
    doc_geral.add_paragraph("[DESCREVER_AQUI_A_DECISAO_FISCAL]")
    doc_geral.add_paragraph("Situação Final:")
    doc_geral.add_paragraph("[DESCREVER_AQUI_O_DESFECHO_FINAL]")
    doc_geral.add_paragraph("--------------------------------------------------")
    doc_geral.add_paragraph("5. Conclusão Geral")
    doc_geral.add_paragraph("Diante do exposto, conclui-se que dos [TOTAL_CONTRIBUINTES] contribuintes notificados no âmbito da DSF nº [NUMERO_DSF], [QUANTIDADE_REGULARIZADOS] promoveram a regularização integral das pendências.")
    doc_geral.add_paragraph("Recomenda-se o encaminhamento para as providências fiscais cabíveis.")
    doc_geral.add_paragraph("[LOCAL_E_DATA]")
    doc_geral.add_paragraph("[ESPAÇO RESERVADO PARA ASSINATURA DIGITAL]")
    doc_geral.add_paragraph("____________________________________________________")
    doc_geral.add_paragraph("[NOME_DO_AUDITOR]")
    doc_geral.add_paragraph("[CARGO_DO_AUDITOR]")
    doc_geral.add_paragraph("Matrícula: [MATRICULA]")
    doc_geral.save(caminho_geral)

    return caminho_individual, caminho_geral


def test_health_check():
    resposta = client.get("/api/health")
    assert resposta.status_code == 200
    payload = resposta.json()
    assert payload["status"] == "ok"
    assert payload["version"] == "1.0.0"


def test_modelos_docx_padrao_apontam_para_storage_config(monkeypatch):
    monkeypatch.delenv("RELATORIO_MODELO_INDIVIDUAL_DOCX", raising=False)
    monkeypatch.delenv("RELATORIO_MODELO_GERAL_DOCX", raising=False)

    diretorio_modelos = Path(relatorio_fiscal.__file__).resolve().parents[2] / "storage" / "_config"

    assert relatorio_fiscal._resolver_caminho_modelo_docx("individual") == (
        diretorio_modelos / "Modelo_Relatorio_Individual.docx"
    )
    assert relatorio_fiscal._resolver_caminho_modelo_docx("geral") == (
        diretorio_modelos / "Modelo_Relatorio_Geral.docx"
    )


def test_endpoints_oracle_diagnostico(monkeypatch):
    estado_mapeamentos = {
        "FONTE_C170": "SEFIN.VW_EFD_C170",
        "FONTE_REG0200": "SEFIN.VW_EFD_REG0200",
    }

    monkeypatch.setattr(
        api,
        "testar_conexao_oracle",
        lambda indice=0: {"status": "ok", "usuario": "TESTE", "banco": "SEFIN", "host": "localhost"},
    )
    monkeypatch.setattr(
        api,
        "listar_objetos_oracle",
        lambda termo=None, limite=200, indice=0: [
            {"owner": "SEFIN", "object_name": "VW_EFD_C170", "object_type": "VIEW"}
        ],
    )
    monkeypatch.setattr(
        api,
        "listar_colunas_objeto_oracle",
        lambda objeto, owner=None, limite=500, indice=0: [
            {
                "owner": "SEFIN",
                "object_name": objeto,
                "column_name": "CNPJ_CONTRIBUINTE",
                "data_type": "VARCHAR2",
                "data_length": 14,
                "data_precision": None,
                "data_scale": None,
                "nullable": "N",
            }
        ],
    )
    monkeypatch.setattr(
        api,
        "detalhar_mapeamento_fontes_oracle",
        lambda: [
            {
                "chave": chave,
                "env_var": f"ORACLE_{chave}",
                "fonte_padrao": chave.replace("FONTE_", "VW_"),
                "fonte_configurada": valor,
                "origem": "persistido",
                "owner": valor.split(".", 1)[0],
                "objeto": valor.split(".", 1)[1],
            }
            for chave, valor in estado_mapeamentos.items()
        ],
    )
    monkeypatch.setattr(
        api,
        "salvar_mapeamento_fontes_oracle",
        lambda payload: estado_mapeamentos.update(
            {chave: valor for chave, valor in payload.items() if valor}
        ),
    )

    resposta_conexao = client.get("/api/oracle/conexao")
    assert resposta_conexao.status_code == 200
    assert resposta_conexao.json()["status"] == "ok"

    resposta_fontes = client.get("/api/oracle/fontes?termo=c170")
    assert resposta_fontes.status_code == 200
    assert resposta_fontes.json()["status"] == "ok"
    assert len(resposta_fontes.json()["objetos"]) == 1

    resposta_colunas = client.get("/api/oracle/colunas/VW_EFD_C170?owner=SEFIN")
    assert resposta_colunas.status_code == 200
    assert resposta_colunas.json()["status"] == "ok"
    assert resposta_colunas.json()["colunas"][0]["column_name"] == "CNPJ_CONTRIBUINTE"

    resposta_mapeamentos = client.get("/api/oracle/mapeamentos")
    assert resposta_mapeamentos.status_code == 200
    assert resposta_mapeamentos.json()["status"] == "ok"
    assert len(resposta_mapeamentos.json()["mapeamentos"]) == 2

    resposta_salvar_mapeamentos = client.put(
        "/api/oracle/mapeamentos",
        json={"mapeamentos": {"FONTE_C170": "SEFIN.VW_EFD_C170_AJUSTADA"}},
    )
    assert resposta_salvar_mapeamentos.status_code == 200
    assert resposta_salvar_mapeamentos.json()["status"] == "ok"

    resposta_validacao = client.get("/api/oracle/mapeamentos/validacao")
    assert resposta_validacao.status_code == 200
    assert resposta_validacao.json()["status"] == "ok"
    assert resposta_validacao.json()["total_ok"] >= 1


def test_oracle_mapeamento_raiz_analisa_sqls_por_diretorio(tmp_path):
    pasta_sql = tmp_path / "sql"
    pasta_sql.mkdir(parents=True, exist_ok=True)
    (pasta_sql / "c170.sql").write_text(
        """
        WITH arquivos_validos AS (
            SELECT r.id AS reg_0000_id
            FROM sped.reg_0000 r
            WHERE r.cnpj = :cnpj
        )
        SELECT c170.cod_item, r200.cod_ncm
        FROM sped.reg_c170 c170
        JOIN arquivos_validos av ON av.reg_0000_id = c170.reg_0000_id
        LEFT JOIN sped.reg_0200 r200 ON r200.reg_0000_id = c170.reg_0000_id
        """,
        encoding="utf-8",
    )

    resposta = client.get(f"/api/oracle/mapeamento-raiz?diretorio={pasta_sql.as_posix()}")

    assert resposta.status_code == 200
    payload = resposta.json()
    assert payload["status"] == "ok"
    assert payload["resumo"]["total_sqls"] == 1
    assert payload["resumo"]["total_fontes_raiz"] == 3
    assert payload["arquivos_sql"][0]["tem_bind_cnpj"] is True
    assert "SPED.REG_C170" in [fonte["fonte_oracle"] for fonte in payload["fontes_raiz"]]

    fonte_c170 = next(fonte for fonte in payload["fontes_raiz"] if fonte["fonte_oracle"] == "SPED.REG_C170")
    assert fonte_c170["camada_bronze"] == "bronze/efd_reg_c170"
    assert fonte_c170["chave_recorte"] == "reg_0000_id"
    assert payload["estrategia_polars"]


def test_status_sistema_expoe_conectividade_oracle_real(tmp_path, monkeypatch):
    monkeypatch.setattr(api, "BASE_DIR", tmp_path)
    monkeypatch.setattr(
        api,
        "obter_configuracao_oracle",
        lambda indice=0: object(),
    )
    monkeypatch.setattr(
        api,
        "_obter_status_conectividade_oracle",
        lambda indice_oracle: (False, "getaddrinfo failed"),
    )
    monkeypatch.setattr(
        api,
        "listar_resumos_configuracoes_oracle",
        lambda indices_extras=None: [
            {
                "indice": 0,
                "host": "oracle.exemplo.gov.br",
                "porta": 1521,
                "servico": "sefindw",
                "configurada": True,
                "erro": None,
            }
        ],
    )

    resposta = client.get("/api/sistema/status")

    assert resposta.status_code == 200
    payload = resposta.json()
    assert payload["oracle_configurada"] is True
    assert payload["oracle_conectada"] is False
    assert payload["oracle_indice_ativo"] == 0
    assert payload["erro_oracle"] == "getaddrinfo failed"
    assert payload["conexoes_oracle"][0]["ativa"] is True


def test_sistema_alvos_lista_cnpjs_com_metadados_operacionais(tmp_path, monkeypatch):
    monkeypatch.setattr(api, "BASE_DIR", tmp_path)
    _preparar_cnpj_teste(tmp_path)
    _escrever_parquet(
        tmp_path / "12345678000190" / "extraidos" / "reg0000.parquet",
        [
            {
                "nome": "EMPRESA TESTE LTDA",
                "ie": "123456789",
                "cpf": "12345678901",
            }
        ],
    )
    _preparar_relatorio_cnpj(
        tmp_path,
        "12345678000190",
        contribuinte="EMPRESA TESTE LTDA",
        dsf="2026000001",
    )

    resposta = client.get("/api/sistema/alvos")

    assert resposta.status_code == 200
    payload = resposta.json()
    assert payload["status"] == "ok"
    assert payload["resumo"]["total_cnpjs"] == 1
    assert payload["resumo"]["total_cpfs_mapeados"] == 1
    assert payload["resumo"]["total_cnpjs_com_pipeline"] == 1
    assert payload["resumo"]["total_cnpjs_com_relatorio"] == 1

    alvo = payload["alvos"][0]
    assert alvo["cnpj"] == "12345678000190"
    assert alvo["contribuinte"] == "EMPRESA TESTE LTDA"
    assert alvo["ie"] == "123456789"
    assert alvo["cpfs_vinculados"] == ["12345678901"]
    assert alvo["possui_relatorio"] is True
    assert alvo["status_pipeline"] == "parcial"
    assert alvo["total_parquets"] >= 1


def test_endpoints_oracle_usam_indice_ativo_quando_query_nao_e_informada(tmp_path, monkeypatch):
    monkeypatch.setattr(api, "BASE_DIR", tmp_path)
    (tmp_path / "_sistema").mkdir(parents=True, exist_ok=True)
    (tmp_path / "_sistema" / "configuracoes.json").write_text(
        json.dumps({"oracle_indice_ativo": 3}),
        encoding="utf-8",
    )

    chamadas: list[int] = []

    monkeypatch.setattr(
        api,
        "testar_conexao_oracle",
        lambda indice=0: chamadas.append(indice) or {"status": "ok", "usuario": "TESTE", "banco": "SEFIN", "host": "localhost"},
    )

    resposta_padrao = client.get("/api/oracle/conexao")
    resposta_override = client.get("/api/oracle/conexao?indice=1")

    assert resposta_padrao.status_code == 200
    assert resposta_override.status_code == 200
    assert chamadas == [3, 1]


def test_consulta_cadastral_retorna_cnpj_encontrado(tmp_path, monkeypatch):
    monkeypatch.setattr(api, "BASE_DIR", tmp_path)
    monkeypatch.setattr(api, "_obter_status_conectividade_oracle", lambda indice_oracle: (True, None))
    monkeypatch.setattr(
        api,
        "consultar_dados_cadastrais_documentos",
        lambda documentos, indice_oracle=0: [
            {
                "status": "ok",
                "tipo_documento": "cnpj",
                "documento_consultado": documentos[0],
                "origem": "oracle",
                "encontrado": True,
                "mensagem": None,
                "registros": [{"documento": documentos[0], "nome": "EMPRESA TESTE LTDA", "ie": "123"}],
            }
        ],
    )

    resposta = client.post("/api/cadastro/consultar", json={"documentos": ["12.345.678/0001-90"]})

    assert resposta.status_code == 200
    payload = resposta.json()
    assert payload["status"] == "ok"
    assert payload["documentos_processados"] == 1
    assert payload["resultados"][0]["tipo_documento"] == "cnpj"
    assert payload["resultados"][0]["documento_consultado"] == "12345678000190"
    assert payload["resultados"][0]["encontrado"] is True


def test_consulta_cadastral_retorna_cpf_encontrado(tmp_path, monkeypatch):
    monkeypatch.setattr(api, "BASE_DIR", tmp_path)
    monkeypatch.setattr(api, "_obter_status_conectividade_oracle", lambda indice_oracle: (True, None))
    monkeypatch.setattr(
        api,
        "consultar_dados_cadastrais_documentos",
        lambda documentos, indice_oracle=0: [
            {
                "status": "ok",
                "tipo_documento": "cpf",
                "documento_consultado": documentos[0],
                "origem": "oracle",
                "encontrado": True,
                "mensagem": None,
                "registros": [{"documento": documentos[0], "nome": "CONTRIBUINTE TESTE", "ie": None}],
            }
        ],
    )

    resposta = client.post("/api/cadastro/consultar", json={"documentos": ["123.456.789-01"]})

    assert resposta.status_code == 200
    payload = resposta.json()
    assert payload["resultados"][0]["tipo_documento"] == "cpf"
    assert payload["resultados"][0]["documento_consultado"] == "12345678901"


def test_consulta_cadastral_retorna_documento_sem_resultado(tmp_path, monkeypatch):
    monkeypatch.setattr(api, "BASE_DIR", tmp_path)
    monkeypatch.setattr(api, "_obter_status_conectividade_oracle", lambda indice_oracle: (True, None))
    monkeypatch.setattr(
        api,
        "consultar_dados_cadastrais_documentos",
        lambda documentos, indice_oracle=0: [
            {
                "status": "ok",
                "tipo_documento": "cnpj",
                "documento_consultado": documentos[0],
                "origem": "oracle",
                "encontrado": False,
                "mensagem": "Nenhum dado cadastral encontrado para o documento informado",
                "registros": [],
            }
        ],
    )

    resposta = client.post("/api/cadastro/consultar", json={"documentos": ["12.345.678/0001-90"]})

    assert resposta.status_code == 200
    payload = resposta.json()
    assert payload["resultados"][0]["encontrado"] is False
    assert payload["resultados"][0]["registros"] == []


def test_consulta_cadastral_trata_lote_misto_com_invalidos_e_duplicados(tmp_path, monkeypatch):
    monkeypatch.setattr(api, "BASE_DIR", tmp_path)
    monkeypatch.setattr(api, "_obter_status_conectividade_oracle", lambda indice_oracle: (True, None))

    chamadas: list[list[str]] = []

    def consultar(documentos, indice_oracle=0):
        chamadas.append(documentos)
        return [
            {
                "status": "ok",
                "tipo_documento": "cnpj" if len(documento) == 14 else "cpf",
                "documento_consultado": documento,
                "origem": "oracle",
                "encontrado": True,
                "mensagem": None,
                "registros": [{"documento": documento}],
            }
            for documento in documentos
        ]

    monkeypatch.setattr(api, "consultar_dados_cadastrais_documentos", consultar)

    resposta = client.post(
        "/api/cadastro/consultar",
        json={"documentos": ["12.345.678/0001-90", "12345678901", "12.345.678/0001-90", "abc"]},
    )

    assert resposta.status_code == 200
    payload = resposta.json()
    assert chamadas == [["12345678000190", "12345678901"]]
    assert len(payload["resultados"]) == 3
    assert sum(1 for item in payload["resultados"] if item["status"] == "invalido") == 1


def test_consulta_cadastral_retorna_erro_quando_oracle_indisponivel(tmp_path, monkeypatch):
    monkeypatch.setattr(api, "BASE_DIR", tmp_path)
    monkeypatch.setattr(api, "_obter_status_conectividade_oracle", lambda indice_oracle: (False, "getaddrinfo failed"))

    resposta = client.post("/api/cadastro/consultar", json={"documentos": ["12.345.678/0001-90"]})

    assert resposta.status_code == 503
    payload = resposta.json()
    assert payload["status"] == "erro"
    assert payload["mensagem"] == "Falha na conexao com o Oracle para consulta cadastral"


def test_fluxo_api_principal(tmp_path, monkeypatch):
    monkeypatch.setattr(api, "BASE_DIR", tmp_path)
    _preparar_cnpj_teste(tmp_path)

    # Configuracoes
    resposta_cfg = client.put(
        "/api/configuracoes",
        json={"reprocessamento_automatico": True, "logs_detalhados": True},
    )
    assert resposta_cfg.status_code == 200
    assert resposta_cfg.json()["status"] == "ok"

    # Execucao pipeline sem extracao Oracle
    resposta_pipeline = client.post(
        "/api/pipeline/executar",
        json={
            "cnpj": "12.345.678/0001-90",
            "consultas": [],
            "executar_extracao": False,
        },
    )
    assert resposta_pipeline.status_code == 200
    payload_pipeline = resposta_pipeline.json()
    assert payload_pipeline["cnpj"] == "12345678000190"
    assert payload_pipeline["status"] in {"concluido", "concluido_com_erros"}

    # Leitura de tabela
    resposta_tabela = client.get("/api/tabelas/12345678000190")
    assert resposta_tabela.status_code == 200
    assert resposta_tabela.json()["status"] == "ok"

    resposta_ler = client.get("/api/tabelas/12345678000190/produtos?pagina=1&por_pagina=10")
    assert resposta_ler.status_code == 200
    assert resposta_ler.json()["total_registros"] >= 0

    # Recria produtos com IDs conhecidos para validar agregacao manual.
    _escrever_parquet(
        tmp_path / "12345678000190" / "parquets" / "produtos.parquet",
        [
            {
                "id_produto": 1,
                "descricao": "Produto A",
                "ncm": "22030000",
                "cest": None,
                "unidade_principal": "UN",
                "qtd_total_nfe": 3,
                "valor_total": 150.0,
                "tipo": "ambos",
            },
            {
                "id_produto": 2,
                "descricao": "Produto B",
                "ncm": "22030000",
                "cest": None,
                "unidade_principal": "UN",
                "qtd_total_nfe": 2,
                "valor_total": 90.0,
                "tipo": "compra",
            },
        ],
    )

    # Agregacao manual
    resposta_agregar = client.post(
        "/api/agregacao/agregar?cnpj=12345678000190",
        json={
            "ids_produtos": ["1", "2"],
            "descricao_padrao": "Grupo Teste",
        },
    )
    assert resposta_agregar.status_code == 200
    assert resposta_agregar.json()["status"] == "ok"

    arquivo_agregacao = tmp_path / "12345678000190" / "edicoes" / "agregacao.json"
    assert arquivo_agregacao.exists()
    dados_agregacao = json.loads(arquivo_agregacao.read_text(encoding="utf-8"))
    assert "Grupo Teste" in dados_agregacao

    # Conversao manual
    resposta_fator = client.put(
        "/api/conversao/fator?cnpj=12345678000190",
        json={
            "id_agrupado": "id_agrupado_1",
            "fator": 2.0,
            "unid_ref": "CX",
        },
    )
    assert resposta_fator.status_code == 200
    assert resposta_fator.json()["status"] == "ok"

    arquivo_fatores = tmp_path / "12345678000190" / "edicoes" / "fatores.json"
    assert arquivo_fatores.exists()

    # Exportacao
    resposta_export = client.get("/api/exportar/12345678000190/produtos?formato=csv")
    assert resposta_export.status_code == 200
    assert resposta_export.headers["content-type"].startswith("text/csv")


def test_api_lista_e_le_camadas_e_expoe_manifesto(tmp_path, monkeypatch):
    monkeypatch.setattr(api, "BASE_DIR", tmp_path)
    _preparar_cnpj_teste(tmp_path)

    resposta_extraidos = client.get("/api/tabelas/12345678000190?camada=extraidos")
    assert resposta_extraidos.status_code == 200
    payload_extraidos = resposta_extraidos.json()
    assert payload_extraidos["camada"] == "extraidos"
    assert any(tabela["nome"] == "reg0000" for tabela in payload_extraidos["tabelas"])

    resposta_silver = client.get("/api/tabelas/12345678000190?camada=silver")
    assert resposta_silver.status_code == 200
    payload_silver = resposta_silver.json()
    assert payload_silver["camada"] == "silver"
    assert any(tabela["nome"] == "item_unidades" for tabela in payload_silver["tabelas"])

    resposta_ler_silver = client.get("/api/tabelas/12345678000190/item_unidades?camada=silver&pagina=1&por_pagina=10")
    assert resposta_ler_silver.status_code == 200
    payload_ler_silver = resposta_ler_silver.json()
    assert payload_ler_silver["camada"] == "silver"
    assert payload_ler_silver["total_registros"] == 1

    resposta_manifesto = client.get("/api/storage/12345678000190/manifesto")
    assert resposta_manifesto.status_code == 200
    manifesto = resposta_manifesto.json()["manifesto"]
    assert manifesto["cnpj"] == "12345678000190"
    assert manifesto["camadas"]["extraidos"]["total_tabelas"] >= 1
    assert manifesto["camadas"]["silver"]["total_tabelas"] >= 1
    assert manifesto["camadas"]["parquets"]["total_tabelas"] >= 1


def test_pipeline_explica_erros_parciais_da_extracao_sem_perder_sucesso_do_pipeline(tmp_path, monkeypatch):
    monkeypatch.setattr(api, "BASE_DIR", tmp_path)
    _preparar_cnpj_teste(tmp_path)

    class _EtapaSucesso:
        tabela = "produtos"
        status = type("StatusFake", (), {"value": "concluida"})()
        mensagem = ""
        duracao_ms = 15
        registros_gerados = 2
        arquivo_saida = "c:/tmp/produtos.parquet"

    class _ResultadoSucesso:
        cnpj = "12345678000190"
        status = "concluido"
        erros: list[str] = []
        etapas = [_EtapaSucesso()]
        duracao_total_ms = 15
        tabelas_geradas = ["produtos"]

    class _OrquestradorSucesso:
        def __init__(self, *_args, **_kwargs):
            pass

        def executar_pipeline_completo(self, tabelas_alvo=None):
            return _ResultadoSucesso()

    monkeypatch.setattr(
        api,
        "extrair_dados_cnpj",
        lambda **_: {
            "cnpj": "12345678000190",
            "diretorio_extraidos": "c:/tmp/extraidos",
            "consultas": {
                "nfce": {"status": "erro", "linhas": 0, "mensagem": "ORA-00942"},
                "nfe": {"status": "ok", "linhas": 10, "arquivo": "nfe.parquet"},
            },
            "erros": ["Falha na consulta nfce: ORA-00942"],
            "status": "concluido_com_erros",
            "total_linhas": 10,
        },
    )
    monkeypatch.setattr(api, "OrquestradorPipeline", _OrquestradorSucesso)
    monkeypatch.setattr(api, "_obter_status_conectividade_oracle", lambda indice_oracle: (True, None))

    resposta = client.post(
        "/api/pipeline/executar",
        json={
            "cnpj": "12.345.678/0001-90",
            "consultas": ["nfce", "nfe"],
            "executar_extracao": True,
        },
    )

    assert resposta.status_code == 200
    payload = resposta.json()
    assert payload["status"] == "concluido_com_erros"
    assert payload["erros"] == []
    assert payload["erros_pipeline"] == []
    assert payload["erros_extracao"] == ["Falha na consulta nfce: ORA-00942"]
    assert payload["erros_total"] == ["Falha na consulta nfce: ORA-00942"]
    assert payload["extracao"]["status"] == "concluido_com_erros"
    assert payload["extracao"]["consultas"]["nfce"]["status"] == "erro"
    assert payload["tabelas_geradas"] == ["produtos"]


def test_pipeline_retorna_erro_json_quando_extracao_oracle_falha(tmp_path, monkeypatch):
    monkeypatch.setattr(api, "BASE_DIR", tmp_path)

    chamadas = {"orquestrador_instanciado": False}

    def _falhar_extracao(**_: object) -> dict:
        raise RuntimeError("getaddrinfo failed")

    class _OrquestradorNaoDeveExecutar:
        def __init__(self, *_args, **_kwargs):
            chamadas["orquestrador_instanciado"] = True
            raise AssertionError("O pipeline nao deve iniciar quando a extracao Oracle falha")

    monkeypatch.setattr(api, "extrair_dados_cnpj", _falhar_extracao)
    monkeypatch.setattr(api, "OrquestradorPipeline", _OrquestradorNaoDeveExecutar)
    monkeypatch.setattr(api, "_obter_status_conectividade_oracle", lambda indice_oracle: (False, "getaddrinfo failed"))

    resposta = client.post(
        "/api/pipeline/executar",
        json={
            "cnpj": "12.345.678/0001-90",
            "consultas": ["nfe"],
            "executar_extracao": True,
        },
    )

    assert resposta.status_code == 503
    payload = resposta.json()
    assert payload["status"] == "erro"
    assert payload["cnpj"] == "12345678000190"
    assert "Oracle" in payload["mensagem"]
    assert "getaddrinfo failed" in payload["detalhe"]
    assert payload["indice_oracle"] == 0
    assert chamadas["orquestrador_instanciado"] is False


def test_pipeline_repassa_data_limite_nula_para_extracao_oracle(tmp_path, monkeypatch):
    monkeypatch.setattr(api, "BASE_DIR", tmp_path)
    _preparar_cnpj_teste(tmp_path)
    (tmp_path / "_sistema").mkdir(parents=True, exist_ok=True)
    (tmp_path / "_sistema" / "configuracoes.json").write_text(
        json.dumps({"diretorio_consultas_sql": "c:/referencia/sql"}),
        encoding="utf-8",
    )

    chamadas_extracao: list[dict] = []

    class _EtapaSucesso:
        tabela = "produtos"
        status = type("StatusFake", (), {"value": "concluida"})()
        mensagem = ""
        duracao_ms = 10
        registros_gerados = 1
        arquivo_saida = "c:/tmp/produtos.parquet"

    class _ResultadoSucesso:
        cnpj = "12345678000190"
        status = "concluido"
        erros: list[str] = []
        etapas = [_EtapaSucesso()]
        duracao_total_ms = 10
        tabelas_geradas = ["produtos"]

    class _OrquestradorSucesso:
        def __init__(self, *_args, **_kwargs):
            pass

        def executar_pipeline_completo(self, tabelas_alvo=None):
            return _ResultadoSucesso()

    def _extracao_falsa(**kwargs):
        chamadas_extracao.append(kwargs)
        return {
            "cnpj": "12345678000190",
            "diretorio_extraidos": "c:/tmp/extraidos",
            "consultas": {
                "c170": {"status": "ok", "linhas": 1, "arquivo": "c170.parquet"},
            },
            "erros": [],
            "status": "ok",
            "total_linhas": 1,
        }

    monkeypatch.setattr(api, "extrair_dados_cnpj", _extracao_falsa)
    monkeypatch.setattr(api, "OrquestradorPipeline", _OrquestradorSucesso)
    monkeypatch.setattr(api, "_obter_status_conectividade_oracle", lambda indice_oracle: (True, None))

    resposta = client.post(
        "/api/pipeline/executar",
        json={
            "cnpj": "12.345.678/0001-90",
            "consultas": ["c170"],
            "executar_extracao": True,
        },
    )

    assert resposta.status_code == 200
    assert len(chamadas_extracao) == 1
    assert chamadas_extracao[0]["data_limite"] is None
    assert chamadas_extracao[0]["diretorio_consultas"] == "c:/referencia/sql"


def test_consultas_usam_diretorio_sql_configurado(tmp_path, monkeypatch):
    monkeypatch.setattr(api, "BASE_DIR", tmp_path)
    (tmp_path / "_sistema").mkdir(parents=True, exist_ok=True)
    (tmp_path / "_sistema" / "configuracoes.json").write_text(
        json.dumps({"diretorio_consultas_sql": "c:/mapa/sql"}),
        encoding="utf-8",
    )

    chamadas: list[str | None] = []
    monkeypatch.setattr(
        api,
        "listar_consultas_versionadas",
        lambda diretorio_consultas=None: chamadas.append(diretorio_consultas) or ["c170", "reg0200"],
    )

    resposta = client.get("/api/consultas")

    assert resposta.status_code == 200
    payload = resposta.json()
    assert payload["consultas"] == ["c170", "reg0200"]
    assert payload["diretorio_consultas_sql"] == "c:/mapa/sql"
    assert chamadas == ["c:/mapa/sql"]


def test_relatorio_diagnostico_lista_dsfs_inferidas_e_pipeline_local(tmp_path, monkeypatch):
    monkeypatch.setattr(api, "BASE_DIR", tmp_path)

    cnpj_pipeline = "37671507000187"
    diretorio_cnpj = tmp_path / cnpj_pipeline
    for pasta in ["extraidos", "parquets", "edicoes", "exportacoes"]:
        (diretorio_cnpj / pasta).mkdir(parents=True, exist_ok=True)

    for contrato in api.CONTRATOS.values():
        _escrever_parquet(diretorio_cnpj / "parquets" / contrato.saida, [], schema={})

    _preparar_relatorio_cnpj(tmp_path, "19288989000290", contribuinte="EMPRESA A")
    _preparar_relatorio_cnpj(tmp_path, "21418376000190", contribuinte="EMPRESA B")

    monkeypatch.setattr(
        api,
        "diagnosticar_prontidao_relatorios",
        lambda _base_dir: {
            "pronto_pdf": False,
            "dependencias": [
                {"nome": "weasyprint", "instalado": False, "mensagem": "Modulo Python ausente: weasyprint"},
                {"nome": "pypdf", "instalado": False, "mensagem": "Modulo Python ausente: pypdf"},
            ],
            "dependencias_faltantes": ["weasyprint", "pypdf"],
            "modelos_docx": {
                "individual": {"tipo": "individual", "variavel_ambiente": "RELATORIO_MODELO_INDIVIDUAL_DOCX", "caminho_resolvido": "modelo_individual.docx", "existe": True, "pronto": True, "mensagem": ""},
                "geral": {"tipo": "geral", "variavel_ambiente": "RELATORIO_MODELO_GERAL_DOCX", "caminho_resolvido": "modelo_geral.docx", "existe": True, "pronto": True, "mensagem": ""},
            },
            "total_cnpjs_com_relatorio": 2,
            "total_dsfs": 1,
            "cnpjs_com_relatorio": api.rf_listar_cnpjs_com_relatorio(tmp_path),
            "dsfs": list(api.carregar_dsfs_efetivas(tmp_path).values()),
        },
    )

    resposta = client.get("/api/relatorio/diagnostico")

    assert resposta.status_code == 200
    payload = resposta.json()
    assert payload["pronto_pdf"] is False
    assert payload["dependencias_faltantes"] == ["weasyprint", "pypdf"]
    assert payload["dsfs"][0]["numero"] == "20263710400226"
    assert payload["dsfs"][0]["origem"] == "inferida"
    assert payload["pipeline_local"]["cnpj_referencia"] == cnpj_pipeline
    assert payload["pipeline_local"]["completo"] is True
    assert payload["pipeline_local"]["total_tabelas_ok"] == len(api.CONTRATOS)


def test_relatorio_dsf_lista_inferida_sem_dsf_json(tmp_path, monkeypatch):
    monkeypatch.setattr(api, "BASE_DIR", tmp_path)
    _preparar_relatorio_cnpj(tmp_path, "19288989000290", contribuinte="EMPRESA A")
    _preparar_relatorio_cnpj(tmp_path, "21418376000190", contribuinte="EMPRESA B")

    resposta_lista = client.get("/api/relatorio/dsf")
    resposta_item = client.get("/api/relatorio/dsf/20263710400226")

    assert resposta_lista.status_code == 200
    assert resposta_item.status_code == 200
    dsfs = resposta_lista.json()["dsfs"]
    assert len(dsfs) == 1
    assert dsfs[0]["numero"] == "20263710400226"
    assert sorted(dsfs[0]["cnpjs"]) == ["19288989000290", "21418376000190"]
    assert resposta_item.json()["dsf"]["origem"] == "inferida"


def test_lista_cnpjs_com_relatorio_informa_det_invalido(tmp_path, monkeypatch):
    monkeypatch.setattr(api, "BASE_DIR", tmp_path)
    _preparar_relatorio_cnpj(
        tmp_path,
        "32846441000103",
        contribuinte="MEGAFORTT COMERCIO DE ALIMENTOS LTDA",
        pdf_det_invalido=True,
    )

    resposta = client.get("/api/relatorio/listar-cnpjs-com-relatorio")

    assert resposta.status_code == 200
    payload = resposta.json()["cnpjs"]
    assert payload[0]["tem_det"] is False
    assert "DET informado nao encontrado" in payload[0]["aviso_det"]


def test_geracao_pdf_individual_retorna_erro_estruturado_quando_dependencia_falta(tmp_path, monkeypatch):
    monkeypatch.setattr(api, "BASE_DIR", tmp_path)
    _preparar_relatorio_cnpj(tmp_path, "19288989000290", contribuinte="EMPRESA A")
    (tmp_path.parent / "_config").mkdir(parents=True, exist_ok=True)
    (tmp_path.parent / "_config" / "auditor.json").write_text(
        json.dumps({"nome": "Auditor Teste", "matricula": "1"}),
        encoding="utf-8",
    )

    monkeypatch.setattr(
        api,
        "diagnosticar_prontidao_relatorios",
        lambda _base_dir: {
            "pronto_pdf": False,
            "dependencias": [],
            "dependencias_faltantes": ["weasyprint"],
            "modelos_docx": {
                "individual": {"tipo": "individual", "variavel_ambiente": "RELATORIO_MODELO_INDIVIDUAL_DOCX", "caminho_resolvido": "modelo_individual.docx", "existe": True, "pronto": True, "mensagem": ""},
                "geral": {"tipo": "geral", "variavel_ambiente": "RELATORIO_MODELO_GERAL_DOCX", "caminho_resolvido": "modelo_geral.docx", "existe": True, "pronto": True, "mensagem": ""},
            },
            "total_cnpjs_com_relatorio": 1,
            "total_dsfs": 1,
            "cnpjs_com_relatorio": [],
            "dsfs": [],
        },
    )
    monkeypatch.setattr(api, "gerar_pdf_individual", lambda *_args, **_kwargs: (_ for _ in ()).throw(RuntimeError("Modulo Python ausente: weasyprint")))

    resposta = client.post("/api/relatorio/cnpj/19288989000290/gerar-pdf")

    assert resposta.status_code == 503
    payload = resposta.json()
    assert payload["status"] == "erro"
    assert payload["erro"]["codigo"] == "relatorio_pdf_individual_indisponivel"
    assert "weasyprint" in payload["erro"]["detalhe"]
    assert payload["diagnostico"]["dependencias_faltantes"] == ["weasyprint"]


def test_geracao_pdf_individual_funciona_com_gerador_disponivel(tmp_path, monkeypatch):
    monkeypatch.setattr(api, "BASE_DIR", tmp_path)
    _preparar_relatorio_cnpj(tmp_path, "19288989000290", contribuinte="EMPRESA A")
    (tmp_path.parent / "_config").mkdir(parents=True, exist_ok=True)
    (tmp_path.parent / "_config" / "auditor.json").write_text(
        json.dumps({"nome": "Auditor Teste", "matricula": "1"}),
        encoding="utf-8",
    )

    monkeypatch.setattr(
        api,
        "diagnosticar_prontidao_relatorios",
        lambda _base_dir: {
            "pronto_pdf": True,
            "dependencias": [],
            "dependencias_faltantes": [],
            "modelos_docx": {
                "individual": {"tipo": "individual", "variavel_ambiente": "RELATORIO_MODELO_INDIVIDUAL_DOCX", "caminho_resolvido": "modelo_individual.docx", "existe": True, "pronto": True, "mensagem": ""},
                "geral": {"tipo": "geral", "variavel_ambiente": "RELATORIO_MODELO_GERAL_DOCX", "caminho_resolvido": "modelo_geral.docx", "existe": True, "pronto": True, "mensagem": ""},
            },
            "total_cnpjs_com_relatorio": 1,
            "total_dsfs": 1,
            "cnpjs_com_relatorio": [],
            "dsfs": [],
        },
    )

    def _gerar_pdf_fake(_dados, _auditor, output_path, _pdf_det):
        _criar_pdf_fake(Path(output_path))

    monkeypatch.setattr(api, "gerar_pdf_individual", _gerar_pdf_fake)

    resposta = client.post("/api/relatorio/cnpj/19288989000290/gerar-pdf")

    assert resposta.status_code == 200
    assert resposta.headers["content-type"] == "application/pdf"


def test_geracao_docx_individual_preenche_template_a_partir_do_dados_json(tmp_path, monkeypatch):
    from docx import Document

    monkeypatch.setattr(api, "BASE_DIR", tmp_path)
    _preparar_relatorio_cnpj(
        tmp_path,
        "32846441000103",
        contribuinte="MEGAFORTT COMERCIO DE ALIMENTOS LTDA",
        manifestacao="Apresentou contestacao",
        manifestacoes={
            "regularizou_integralmente": False,
            "apresentou_contestacao": True,
            "solicitou_prorrogacao": False,
            "nao_apresentou_manifestacao": False,
        },
    )
    (tmp_path.parent / "_config").mkdir(parents=True, exist_ok=True)
    (tmp_path.parent / "_config" / "auditor.json").write_text(
        json.dumps(
            {
                "nome": "Auditor Teste",
                "cargo": "Auditor Fiscal",
                "matricula": "123",
                "local_data": "Porto Velho, 29 de março de 2026",
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    modelo_individual, modelo_geral = _criar_modelos_docx_teste(tmp_path)
    monkeypatch.setenv("RELATORIO_MODELO_INDIVIDUAL_DOCX", str(modelo_individual))
    monkeypatch.setenv("RELATORIO_MODELO_GERAL_DOCX", str(modelo_geral))

    resposta = client.post("/api/relatorio/cnpj/32846441000103/gerar-docx")

    assert resposta.status_code == 200
    assert (
        resposta.headers["content-type"]
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    documento = Document(BytesIO(resposta.content))
    textos_paragrafos = [paragrafo.text for paragrafo in documento.paragraphs]
    textos_tabelas = [
        celula.text
        for tabela in documento.tables
        for linha in tabela.rows
        for celula in linha.cells
    ]
    texto_completo = "\n".join(textos_paragrafos + textos_tabelas)

    assert "MEGAFORTT COMERCIO DE ALIMENTOS LTDA" in texto_completo
    assert "Contato realizado por e-mail." in texto_completo
    assert "Encaminhar para acao fiscal." in texto_completo
    assert "Pendencia mantida." in texto_completo
    assert "Auditor Teste" in texto_completo
    assert "[NOME_DO_CONTRIBUINTE]" not in texto_completo
    assert any(
        texto.startswith("(X) Apresentou contestação") or texto.startswith("[X] Apresentou contestação")
        for texto in textos_paragrafos + textos_tabelas
    )


def test_geracao_relatorio_geral_resolve_dsfs_inferidas(tmp_path, monkeypatch):
    monkeypatch.setattr(api, "BASE_DIR", tmp_path)
    _preparar_relatorio_cnpj(tmp_path, "19288989000290", contribuinte="EMPRESA A")
    _preparar_relatorio_cnpj(tmp_path, "21418376000190", contribuinte="EMPRESA B")
    (tmp_path.parent / "_config").mkdir(parents=True, exist_ok=True)
    (tmp_path.parent / "_config" / "auditor.json").write_text(
        json.dumps({"nome": "Auditor Teste", "matricula": "1"}),
        encoding="utf-8",
    )

    monkeypatch.setattr(
        api,
        "diagnosticar_prontidao_relatorios",
        lambda _base_dir: {
            "pronto_pdf": True,
            "dependencias": [],
            "dependencias_faltantes": [],
            "modelos_docx": {
                "individual": {"tipo": "individual", "variavel_ambiente": "RELATORIO_MODELO_INDIVIDUAL_DOCX", "caminho_resolvido": "modelo_individual.docx", "existe": True, "pronto": True, "mensagem": ""},
                "geral": {"tipo": "geral", "variavel_ambiente": "RELATORIO_MODELO_GERAL_DOCX", "caminho_resolvido": "modelo_geral.docx", "existe": True, "pronto": True, "mensagem": ""},
            },
            "total_cnpjs_com_relatorio": 2,
            "total_dsfs": 1,
            "cnpjs_com_relatorio": [],
            "dsfs": list(api.carregar_dsfs_efetivas(tmp_path).values()),
        },
    )

    def _gerar_pdf_fake(_empresas, _auditor, output_path, _incluir_dets):
        _criar_pdf_fake(Path(output_path))

    monkeypatch.setattr(api, "gerar_pdf_geral", _gerar_pdf_fake)

    resposta = client.post(
        "/api/relatorio/gerar-geral",
        json={"dsf": "20263710400226", "cnpjs": [], "incluir_dets": True},
    )

    assert resposta.status_code == 200
    assert resposta.headers["content-type"] == "application/pdf"


def test_relatorio_cnpj_normaliza_manifestacoes_legadas_e_seleciona_dets_locais(tmp_path, monkeypatch):
    monkeypatch.setattr(api, "BASE_DIR", tmp_path)
    _preparar_relatorio_cnpj(
        tmp_path,
        "23722199000167",
        contribuinte="CARLOS LTDA",
        manifestacao="Solicitou prorrogacao de prazo",
    )
    _criar_pdf_textual(tmp_path / "23722199000167" / "relatorio" / "det_um.pdf", "DET UM")
    _criar_pdf_textual(tmp_path / "23722199000167" / "relatorio" / "det_dois.pdf", "DET DOIS")

    resposta_relatorio = client.get("/api/relatorio/cnpj/23722199000167")
    resposta_dets = client.get("/api/relatorio/cnpj/23722199000167/listar-dets")

    assert resposta_relatorio.status_code == 200
    dados = resposta_relatorio.json()["dados"]
    assert dados["manifestacoes"]["solicitou_prorrogacao"] is True
    assert dados["manifestacoes"]["nao_apresentou_manifestacao"] is False
    assert sorted(dados["arquivos_notificacao_incluidos"]) == ["det_dois.pdf", "det_um.pdf"]

    assert resposta_dets.status_code == 200
    dets = resposta_dets.json()["dets"]
    assert all(item["selecionado"] is True for item in dets)


def test_relatorio_cnpj_reconhece_pdf_notificacao_local_pelo_numero_da_notificacao(tmp_path, monkeypatch):
    monkeypatch.setattr(api, "BASE_DIR", tmp_path)
    _preparar_relatorio_cnpj(
        tmp_path,
        "30296058000195",
        contribuinte="M. N. PINHEIRO",
    )
    _criar_pdf_textual(
        tmp_path / "30296058000195" / "relatorio" / "NOTIFICAÇÃONºDET-001MN.pdf",
        "DET NOTIFICACAO",
    )

    resposta_relatorio = client.get("/api/relatorio/cnpj/30296058000195")
    resposta_dets = client.get("/api/relatorio/cnpj/30296058000195/listar-dets")

    assert resposta_relatorio.status_code == 200
    dados = resposta_relatorio.json()["dados"]
    assert dados["tem_det"] is True
    assert dados["arquivos_notificacao_incluidos"] == ["NOTIFICAÇÃONºDET-001MN.pdf"]

    assert resposta_dets.status_code == 200
    dets = resposta_dets.json()["dets"]
    assert len(dets) == 1
    assert dets[0]["nome"] == "NOTIFICAÇÃONºDET-001MN.pdf"
    assert dets[0]["selecionado"] is True


def test_geracao_pdf_individual_real_usa_checklist_e_apenas_dets_selecionados(tmp_path, monkeypatch):
    from pypdf import PdfReader

    monkeypatch.setattr(api, "BASE_DIR", tmp_path)
    _instalar_renderizador_pdf_teste(monkeypatch)
    modelo_individual, modelo_geral = _criar_modelos_docx_teste(tmp_path)
    monkeypatch.setenv("RELATORIO_MODELO_INDIVIDUAL_DOCX", str(modelo_individual))
    monkeypatch.setenv("RELATORIO_MODELO_GERAL_DOCX", str(modelo_geral))

    _preparar_relatorio_cnpj(
        tmp_path,
        "23722199000167",
        contribuinte="CARLOS LTDA",
        manifestacoes={
            "regularizou_integralmente": True,
            "apresentou_contestacao": False,
            "solicitou_prorrogacao": True,
            "nao_apresentou_manifestacao": False,
        },
        arquivos_notificacao_incluidos=["det_um.pdf"],
    )
    _criar_pdf_textual(tmp_path / "23722199000167" / "relatorio" / "det_um.pdf", "DET UM")
    _criar_pdf_textual(tmp_path / "23722199000167" / "relatorio" / "det_dois.pdf", "DET DOIS")
    (tmp_path.parent / "_config").mkdir(parents=True, exist_ok=True)
    (tmp_path.parent / "_config" / "auditor.json").write_text(
        json.dumps({"nome": "Auditor Teste", "matricula": "1", "cargo": "Auditor", "local_data": "Porto Velho, 28 de marco de 2026"}),
        encoding="utf-8",
    )

    resposta = client.post("/api/relatorio/cnpj/23722199000167/gerar-pdf")

    assert resposta.status_code == 200
    leitor = PdfReader(BytesIO(resposta.content))
    texto_total = "\n".join(pagina.extract_text() or "" for pagina in leitor.pages)
    assert "Regularizou integralmente as pend" in texto_total
    assert "Solicitou prorroga" in texto_total
    assert "DET UM" in texto_total
    assert "DET DOIS" not in texto_total


def test_geracao_pdf_geral_real_conta_regularizados_e_respeita_selecao_de_dets(tmp_path, monkeypatch):
    from pypdf import PdfReader

    monkeypatch.setattr(api, "BASE_DIR", tmp_path)
    _instalar_renderizador_pdf_teste(monkeypatch)
    modelo_individual, modelo_geral = _criar_modelos_docx_teste(tmp_path)
    monkeypatch.setenv("RELATORIO_MODELO_INDIVIDUAL_DOCX", str(modelo_individual))
    monkeypatch.setenv("RELATORIO_MODELO_GERAL_DOCX", str(modelo_geral))

    _preparar_relatorio_cnpj(
        tmp_path,
        "19288989000290",
        contribuinte="EMPRESA A",
        manifestacoes={
            "regularizou_integralmente": True,
            "apresentou_contestacao": False,
            "solicitou_prorrogacao": False,
            "nao_apresentou_manifestacao": False,
        },
        arquivos_notificacao_incluidos=["det_a.pdf"],
    )
    _preparar_relatorio_cnpj(
        tmp_path,
        "21418376000190",
        contribuinte="EMPRESA B",
        manifestacoes={
            "regularizou_integralmente": False,
            "apresentou_contestacao": True,
            "solicitou_prorrogacao": False,
            "nao_apresentou_manifestacao": False,
        },
        arquivos_notificacao_incluidos=[],
    )
    _criar_pdf_textual(tmp_path / "19288989000290" / "relatorio" / "det_a.pdf", "DET A")
    _criar_pdf_textual(tmp_path / "21418376000190" / "relatorio" / "det_b.pdf", "DET B")
    (tmp_path.parent / "_config").mkdir(parents=True, exist_ok=True)
    (tmp_path.parent / "_config" / "auditor.json").write_text(
        json.dumps({"nome": "Auditor Teste", "matricula": "1", "cargo": "Auditor", "local_data": "Porto Velho, 28 de marco de 2026"}),
        encoding="utf-8",
    )

    resposta = client.post(
        "/api/relatorio/gerar-geral",
        json={"dsf": "20263710400226", "cnpjs": [], "incluir_dets": True},
    )

    assert resposta.status_code == 200
    leitor = PdfReader(BytesIO(resposta.content))
    texto_total = "\n".join(pagina.extract_text() or "" for pagina in leitor.pages)
    texto_linear = " ".join(texto_total.split())
    assert "1 promoveram a regularização integral" in texto_linear
    assert "DET A" in texto_total
    assert "DET B" not in texto_total
