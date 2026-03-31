"""Modulo de geracao de Relatorios Fiscais Conclusivos - Fisconforme Nao Cumprido.

Gera PDFs profissionais por CNPJ e consolida em relatorio geral.
Os PDFs DET sao anexados intactos ao final de cada relatorio.
A logo do Governo do Estado de Rondonia e embutida via base64.
"""

from __future__ import annotations

import base64
import importlib.util
import json
import logging
import os
import re
import shutil
import tempfile
import unicodedata
from pathlib import Path
from typing import Any, Optional

logger = logging.getLogger(__name__)

CHAVES_MANIFESTACAO = (
    "regularizou_integralmente",
    "apresentou_contestacao",
    "solicitou_prorrogacao",
    "nao_apresentou_manifestacao",
)

ROTULOS_MANIFESTACAO = {
    "regularizou_integralmente": "Regularizou integralmente as pendencias",
    "apresentou_contestacao": "Apresentou contestacao",
    "solicitou_prorrogacao": "Solicitou prorrogacao de prazo",
    "nao_apresentou_manifestacao": "Nao apresentou manifestacao",
}

VARIAVEIS_MODELO_DOCX = {
    "individual": "RELATORIO_MODELO_INDIVIDUAL_DOCX",
    "geral": "RELATORIO_MODELO_GERAL_DOCX",
}

DIRETORIO_MODELOS_RELATORIO = Path(__file__).resolve().parents[2] / "storage" / "_config"

CAMINHOS_MODELO_DOCX_PADRAO = {
    "individual": DIRETORIO_MODELOS_RELATORIO / "Modelo_Relatorio_Individual.docx",
    "geral": DIRETORIO_MODELOS_RELATORIO / "Modelo_Relatorio_Geral.docx",
}

_SENTINELA_SELECAO_AUSENTE = object()


# ============================================================
# LOGO - carregada do arquivo logo_cabecalho.png ao lado deste modulo
# ============================================================
_LOGO_BASE64: str | None = None


def _carregar_logo_base64() -> str:
    """Carrega a logo do cabecalho como base64 a partir do arquivo PNG."""
    global _LOGO_BASE64
    if _LOGO_BASE64 is not None:
        return _LOGO_BASE64

    caminho_logo = Path(__file__).resolve().parent / "logo_cabecalho.png"
    if caminho_logo.exists():
        with caminho_logo.open("rb") as f:
            _LOGO_BASE64 = base64.b64encode(f.read()).decode()
    else:
        _LOGO_BASE64 = ""
        logger.warning("Logo nao encontrada em %s", caminho_logo)

    return _LOGO_BASE64


def _normalizar_texto_busca(valor: Any) -> str:
    """Normaliza texto para comparacoes livres sem acentos."""
    texto = str(valor or "").strip().lower()
    texto_sem_acento = unicodedata.normalize("NFKD", texto)
    return "".join(
        caractere
        for caractere in texto_sem_acento
        if not unicodedata.combining(caractere)
    )


def _importar_weasyprint_html() -> Any:
    """Importa o renderizador HTML->PDF apenas quando necessario."""
    try:
        from weasyprint import HTML
    except ModuleNotFoundError as erro:
        raise RuntimeError(
            "Modulo Python ausente: weasyprint"
        ) from erro
    except Exception as erro:  # noqa: BLE001
        raise RuntimeError(str(erro)) from erro
    return HTML


def _importar_pdf_writer() -> Any:
    """Importa o mesclador de PDFs apenas quando necessario."""
    try:
        from pypdf import PdfWriter
    except ModuleNotFoundError as erro:
        raise RuntimeError(
            "Modulo Python ausente: pypdf"
        ) from erro
    return PdfWriter


def _diagnosticar_dependencia_weasyprint() -> dict[str, Any]:
    """Valida importacao efetiva do WeasyPrint."""
    if importlib.util.find_spec("weasyprint") is None:
        return {"nome": "weasyprint", "instalado": False, "mensagem": "Modulo Python ausente: weasyprint"}
    try:
        _importar_weasyprint_html()
    except Exception as erro:  # noqa: BLE001
        return {"nome": "weasyprint", "instalado": False, "mensagem": str(erro)}
    return {"nome": "weasyprint", "instalado": True, "mensagem": ""}


def _diagnosticar_dependencia_reportlab() -> dict[str, Any]:
    """Valida o renderizador alternativo puro Python."""
    if importlib.util.find_spec("reportlab") is None:
        return {"nome": "reportlab", "instalado": False, "mensagem": "Modulo Python ausente: reportlab"}
    try:
        from reportlab.pdfgen import canvas  # noqa: F401
    except Exception as erro:  # noqa: BLE001
        return {"nome": "reportlab", "instalado": False, "mensagem": str(erro)}
    return {"nome": "reportlab", "instalado": True, "mensagem": ""}


def _diagnosticar_dependencia_pypdf() -> dict[str, Any]:
    """Valida o mesclador de PDFs."""
    if importlib.util.find_spec("pypdf") is None:
        return {"nome": "pypdf", "instalado": False, "mensagem": "Modulo Python ausente: pypdf"}
    try:
        _importar_pdf_writer()
    except Exception as erro:  # noqa: BLE001
        return {"nome": "pypdf", "instalado": False, "mensagem": str(erro)}
    return {"nome": "pypdf", "instalado": True, "mensagem": ""}


def _diagnosticar_dependencia_python_docx() -> dict[str, Any]:
    """Valida leitura dos modelos Word."""
    if importlib.util.find_spec("docx") is None:
        return {"nome": "python-docx", "instalado": False, "mensagem": "Modulo Python ausente: python-docx"}
    try:
        from docx import Document  # noqa: F401
    except Exception as erro:  # noqa: BLE001
        return {"nome": "python-docx", "instalado": False, "mensagem": str(erro)}
    return {"nome": "python-docx", "instalado": True, "mensagem": ""}


def _resolver_caminho_modelo_docx(tipo: str) -> Path:
    """Resolve o caminho configurado do modelo Word."""
    variavel = VARIAVEIS_MODELO_DOCX[tipo]
    caminho_env = os.getenv(variavel, "").strip()
    if caminho_env:
        return Path(caminho_env).expanduser()
    return CAMINHOS_MODELO_DOCX_PADRAO[tipo]


def diagnosticar_modelos_docx() -> dict[str, dict[str, Any]]:
    """Resume localizacao e disponibilidade dos modelos Word."""
    dependencia_docx = _diagnosticar_dependencia_python_docx()
    modelos: dict[str, dict[str, Any]] = {}
    for tipo in ("individual", "geral"):
        caminho = _resolver_caminho_modelo_docx(tipo)
        existe = dependencia_docx["instalado"] and caminho.exists() and caminho.is_file()
        mensagem = "" if existe else (
            dependencia_docx["mensagem"]
            if not dependencia_docx["instalado"]
            else f"Modelo Word {tipo} nao encontrado: {caminho}"
        )
        modelos[tipo] = {
            "tipo": tipo,
            "variavel_ambiente": VARIAVEIS_MODELO_DOCX[tipo],
            "caminho_resolvido": str(caminho),
            "existe": existe,
            "pronto": bool(dependencia_docx["instalado"] and existe),
            "mensagem": mensagem,
        }
    return modelos


def _importar_documento_docx() -> Any:
    """Importa o construtor de documentos DOCX apenas quando necessario."""
    try:
        from docx import Document
    except ModuleNotFoundError as erro:
        raise RuntimeError("Modulo Python ausente: python-docx") from erro
    return Document


def _iterar_paragrafos_das_tabelas(tabelas: Any) -> Any:
    """Percorre todos os paragrafos contidos em tabelas e subtabelas."""
    for tabela in tabelas:
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs:
                    yield paragrafo
                yield from _iterar_paragrafos_das_tabelas(celula.tables)


def _iterar_paragrafos_documento(documento: Any) -> Any:
    """Percorre paragrafos do corpo principal e das tabelas do documento."""
    for paragrafo in documento.paragraphs:
        yield paragrafo
    yield from _iterar_paragrafos_das_tabelas(documento.tables)


def _substituir_texto_do_paragrafo(paragrafo: Any, novo_texto: str) -> None:
    """Substitui o conteudo textual preservando o estilo basico do primeiro run."""
    run_referencia = paragrafo.runs[0] if paragrafo.runs else None
    estilo_fonte = {
        "bold": run_referencia.bold if run_referencia else None,
        "italic": run_referencia.italic if run_referencia else None,
        "underline": run_referencia.underline if run_referencia else None,
        "name": run_referencia.font.name if run_referencia else None,
        "size": run_referencia.font.size if run_referencia else None,
    }

    for run in list(paragrafo.runs):
        paragrafo._element.remove(run._element)

    novo_run = paragrafo.add_run(novo_texto)
    if estilo_fonte["bold"] is not None:
        novo_run.bold = estilo_fonte["bold"]
    if estilo_fonte["italic"] is not None:
        novo_run.italic = estilo_fonte["italic"]
    if estilo_fonte["underline"] is not None:
        novo_run.underline = estilo_fonte["underline"]
    if estilo_fonte["name"]:
        novo_run.font.name = estilo_fonte["name"]
    if estilo_fonte["size"] is not None:
        novo_run.font.size = estilo_fonte["size"]


def _aplicar_substituicoes_em_documento(documento: Any, substituicoes: dict[str, str]) -> None:
    """Aplica substituicoes textuais simples em todo o documento DOCX."""
    for paragrafo in _iterar_paragrafos_documento(documento):
        texto_atual = paragrafo.text
        texto_substituido = texto_atual
        for marcador, valor in substituicoes.items():
            texto_substituido = texto_substituido.replace(marcador, valor)
        if texto_substituido != texto_atual:
            _substituir_texto_do_paragrafo(paragrafo, texto_substituido)


def _montar_substituicoes_relatorio_individual(
    empresa: dict[str, Any],
    auditor: dict[str, str],
) -> dict[str, str]:
    """Monta o mapa de placeholders do relatorio individual por CNPJ."""
    checklist = _determinar_checklist_manifestacao(empresa)

    def marcar_item(texto_base: str, ativo: bool, prefixo: str) -> str:
        marcador = "X" if ativo else " "
        if prefixo == "colchetes":
            return f"[{marcador}] {texto_base}"
        return f"({marcador}) {texto_base}"

    substituicoes = {
        "[NOME_DO_CONTRIBUINTE]": str(empresa.get("contribuinte", "")),
        "[CNPJ]": str(empresa.get("cnpj", "")),
        "[INSCRICAO_ESTADUAL]": str(empresa.get("ie", "")),
        "[NUMERO_DSF]": str(empresa.get("dsf", "")),
        "[NUMERO_DET]": str(empresa.get("notificacao_det", "")),
        "[DESCREVER_AQUI_OS_CONTATOS_REALIZADOS]": str(empresa.get("contatos_realizados", "")),
        "[DESCREVER_AQUI_A_DECISAO_FISCAL]": str(empresa.get("decisao_fiscal", "")),
        "[DESCREVER_AQUI_O_DESFECHO_FINAL]": str(empresa.get("desfecho", "")),
        "[LOCAL_E_DATA]": str(auditor.get("local_data", "")),
        "[ESPACO_RESERVADO_PARA_ASSINATURA_DIGITAL]": "Espaco reservado para assinatura digital",
        "[ESPAÇO RESERVADO PARA ASSINATURA DIGITAL]": "Espaco reservado para assinatura digital",
        "[NOME_DO_AUDITOR]": str(auditor.get("nome", "")),
        "[CARGO_DO_AUDITOR]": str(auditor.get("cargo", "")),
        "[MATRICULA]": str(auditor.get("matricula", "")),
    }

    itens_manifestacao = [
        (
            ["Regularizou integralmente as pendencias", "Regularizou integralmente as pendências"],
            checklist["regularizou"],
        ),
        (
            ["Apresentou contestacao", "Apresentou contestação"],
            checklist["contestou"],
        ),
        (
            ["Solicitou prorrogacao de prazo", "Solicitou prorrogação de prazo"],
            checklist["prorrogacao"],
        ),
        (
            ["Nao apresentou manifestacao", "Não apresentou manifestação"],
            checklist["nao_manifestou"],
        ),
    ]
    for variacoes_texto, ativo in itens_manifestacao:
        texto_saida = variacoes_texto[-1]
        for texto_base in variacoes_texto:
            substituicoes[f"[ ] {texto_base}"] = marcar_item(texto_saida, ativo, "colchetes")
            substituicoes[f"[X] {texto_base}"] = marcar_item(texto_saida, ativo, "colchetes")
            substituicoes[f"( ) {texto_base}"] = marcar_item(texto_saida, ativo, "parenteses")
            substituicoes[f"(X) {texto_base}"] = marcar_item(texto_saida, ativo, "parenteses")

    return substituicoes


def gerar_docx_individual(
    empresa: dict[str, Any],
    auditor: dict[str, str],
    output_path: str,
) -> str:
    """Gera o DOCX individual preenchendo o modelo Word com dados do CNPJ."""
    caminho_modelo = _resolver_caminho_modelo_docx("individual")
    if not caminho_modelo.exists():
        raise RuntimeError(f"Modelo Word individual nao encontrado: {caminho_modelo}")

    Document = _importar_documento_docx()
    documento = Document(str(caminho_modelo))
    substituicoes = _montar_substituicoes_relatorio_individual(empresa, auditor)
    _aplicar_substituicoes_em_documento(documento, substituicoes)
    documento.save(output_path)
    return output_path


# ============================================================
# CSS COMPARTILHADO
# ============================================================
CSS_BASE = """
    @page {
        size: A4;
        margin: 1.8cm 2.2cm 2.2cm 2.2cm;
        @bottom-center {
            content: "SEFIN/CRE/GEFIS - Gerencia de Fiscalizacao | Av. Farquar, n. 2986 - Palacio Rio Madeira - Porto Velho/RO";
            font-size: 7pt;
            color: #888;
        }
        @bottom-right {
            content: "Pagina " counter(page);
            font-size: 8pt;
            color: #888;
        }
    }

    * { margin: 0; padding: 0; box-sizing: border-box; }

    body {
        font-family: 'Noto Sans', 'DejaVu Sans', Arial, Helvetica, sans-serif;
        font-size: 10.5pt;
        line-height: 1.45;
        color: #222;
    }

    .header {
        text-align: center;
        border-top: 2px solid #154360;
        border-bottom: 2px solid #154360;
        padding: 8px 0;
        margin-bottom: 14px;
    }
    .header img.logo {
        display: block;
        margin: 0 auto;
        max-width: 420px;
        height: auto;
    }
    .header-texto {
        display: none;
    }
    .header-fallback h1 {
        font-size: 13pt; color: #154360; margin: 0;
        letter-spacing: 1.5px; text-transform: uppercase;
    }
    .header-fallback h2 {
        font-size: 10.5pt; color: #2c3e50; margin: 2px 0; font-weight: 400;
    }

    .titulo-relatorio {
        background: #154360; color: #fff; text-align: center;
        padding: 7px 15px; margin: 10px 0;
        font-size: 12.5pt; font-weight: 700;
        letter-spacing: 2.5px; text-transform: uppercase;
    }

    .dados-contribuinte {
        background-color: #f4f6f7;
        border-left: 3.5px solid #154360;
        padding: 8px 12px; margin: 10px 0;
    }
    .dados-contribuinte table { width: 100%; border-collapse: collapse; }
    .dados-contribuinte td { padding: 2px 6px; vertical-align: top; font-size: 10pt; }
    .dados-contribuinte td.label {
        font-weight: 700; color: #154360; width: 145px; white-space: nowrap;
    }

    .secao { margin: 10px 0 5px 0; }
    .secao-titulo {
        background-color: #d6eaf8; border-left: 3.5px solid #2980b9;
        padding: 4px 10px; font-size: 11pt; font-weight: 700;
        color: #154360; margin-bottom: 5px;
    }
    .secao-corpo { text-align: justify; margin: 4px 0 8px 0; padding: 0 2px; }
    .secao-corpo p { margin: 4px 0; }

    .lista-procedimentos { margin: 4px 0; padding-left: 18px; }
    .lista-procedimentos li { margin-bottom: 2px; font-size: 10pt; }
    .lista-procedimentos li::marker { color: #2980b9; }

    .manifestacao-box {
        border: 1px solid #d5dbdb; padding: 6px 10px;
        margin: 5px 0; background-color: #fdfefe; font-size: 10pt;
    }
    .manifestacao-item { margin: 2px 0; }
    .check {
        display: inline-block; width: 13px; height: 13px;
        border: 1.2px solid #333; text-align: center;
        line-height: 10px; margin-right: 6px; font-size: 10pt;
        vertical-align: middle; color: #333;
    }
    .check.marcado { background-color: transparent; color: #333; font-weight: 700; }

    .decisao-box {
        background-color: #fef9e7; border: 1px solid #f9e79f;
        border-left: 3.5px solid #f39c12;
        padding: 7px 10px; margin: 5px 0; font-weight: 700; font-size: 10pt;
    }
    .desfecho-box {
        background-color: #fdedec; border: 1px solid #f5b7b1;
        border-left: 3.5px solid #e74c3c;
        padding: 7px 10px; margin: 5px 0; font-size: 10pt;
    }

    .assinatura {
        text-align: center;
        margin-top: 25px;
        padding-top: 8px;
    }
    .assinatura .local-data {
        text-align: right; margin-bottom: 10px;
        font-style: italic; font-size: 10pt;
    }
    .assinatura .espaco-assinatura {
        height: 80px;
        margin: 0 auto;
        max-width: 350px;
        border: 1px dashed #bbb;
        border-radius: 6px;
        display: flex;
        align-items: center;
        justify-content: center;
        color: #bbb;
        font-size: 8pt;
        font-style: italic;
    }
    .assinatura .linha {
        display: inline-block;
        border-top: 1px solid #444;
        min-width: 300px;
        padding-top: 4px;
        margin-top: 5px;
    }
    .assinatura .nome { font-weight: 700; font-size: 11pt; color: #154360; }
    .assinatura .cargo, .assinatura .matricula { font-size: 9.5pt; color: #555; }

    .subtitulo-geral {
        font-size: 10pt; text-align: center;
        color: #555; margin-bottom: 10px;
    }
    .tabela-resumo {
        width: 100%; border-collapse: collapse; margin: 10px 0;
        font-size: 9pt;
    }
    .tabela-resumo th {
        background-color: #154360; color: #fff;
        padding: 6px 8px; text-align: left; font-weight: 600;
        border: 1px solid #154360;
    }
    .tabela-resumo td {
        padding: 5px 8px; border: 1px solid #ddd;
        vertical-align: top;
    }
    .tabela-resumo tr:nth-child(even) { background-color: #f8f9fa; }

    .numero-ordem {
        display: inline-block;
        background-color: #154360; color: #fff;
        width: 22px; height: 22px; border-radius: 50%;
        text-align: center; line-height: 22px;
        font-size: 10pt; font-weight: 700;
        margin-right: 8px; vertical-align: middle;
    }

    .separador-cnpj { page-break-before: always; }

    .resumo-individual {
        border: 1px solid #ddd;
        border-radius: 4px;
        padding: 10px 14px;
        margin: 8px 0;
        background-color: #fafafa;
    }
    .resumo-individual p { margin: 3px 0; font-size: 10pt; }
    .resumo-individual .campo-label { font-weight: 700; color: #154360; }
"""


def _html_cabecalho() -> str:
    """Gera HTML do cabecalho com logo ou fallback textual."""
    logo_b64 = _carregar_logo_base64()
    if logo_b64:
        return f"""
<div class="header">
    <img class="logo" src="data:image/png;base64,{logo_b64}" alt="Governo do Estado de Rondonia" />
</div>"""
    else:
        return """
<div class="header header-fallback">
    <h1>Governo do Estado de Rondonia</h1>
    <h2>Secretaria de Estado de Financas</h2>
    <h2>Coordenadoria da Receita Estadual</h2>
</div>"""


def _html_assinatura(auditor: dict[str, str]) -> str:
    return f"""
<div class="assinatura">
    <div class="local-data">{auditor.get('local_data', '')}</div>
    <div class="espaco-assinatura">Espaco reservado para assinatura digital</div>
    <div class="linha">
        <div class="nome">{auditor.get('nome', '')}</div>
        <div class="cargo">{auditor.get('cargo', '')}</div>
        <div class="matricula">Matricula: {auditor.get('matricula', '')}</div>
    </div>
</div>"""


def _criar_manifestacoes_vazias() -> dict[str, bool]:
    """Cria o objeto de manifestacoes sem nenhuma opcao marcada."""
    return {chave: False for chave in CHAVES_MANIFESTACAO}


def criar_manifestacoes_padrao() -> dict[str, bool]:
    """Retorna o estado padrao usado para novos relatorios."""
    manifestacoes = _criar_manifestacoes_vazias()
    manifestacoes["nao_apresentou_manifestacao"] = True
    return manifestacoes


def inferir_manifestacoes_a_partir_texto(manifestacao: str) -> dict[str, bool]:
    """Infere o checklist a partir do texto legado salvo em `manifestacao`."""
    texto = _normalizar_texto_busca(manifestacao)
    manifestacoes = _criar_manifestacoes_vazias()
    if "regularizou" in texto:
        manifestacoes["regularizou_integralmente"] = True
    if "contest" in texto:
        manifestacoes["apresentou_contestacao"] = True
    if "prorrog" in texto:
        manifestacoes["solicitou_prorrogacao"] = True
    if "nao apresentou manifest" in texto or ("nao apresentou" in texto and "manifest" in texto):
        manifestacoes["nao_apresentou_manifestacao"] = True
    if not any(manifestacoes.values()) and not texto:
        manifestacoes["nao_apresentou_manifestacao"] = True
    return manifestacoes


def normalizar_manifestacoes(
    manifestacoes_informadas: Optional[dict[str, Any]],
    manifestacao_legada: str,
) -> dict[str, bool]:
    """Normaliza o checklist de manifestacoes preservando compatibilidade."""
    manifestacoes_inferidas = inferir_manifestacoes_a_partir_texto(manifestacao_legada)
    if not isinstance(manifestacoes_informadas, dict):
        return manifestacoes_inferidas

    manifestacoes_normalizadas = _criar_manifestacoes_vazias()
    for chave in CHAVES_MANIFESTACAO:
        if chave in manifestacoes_informadas:
            manifestacoes_normalizadas[chave] = bool(manifestacoes_informadas[chave])
        else:
            manifestacoes_normalizadas[chave] = manifestacoes_inferidas[chave]
    return manifestacoes_normalizadas


def resumir_manifestacoes(manifestacoes: dict[str, bool]) -> str:
    """Resume as opcoes marcadas para uso em listagens e consolidado."""
    marcadas = [
        ROTULOS_MANIFESTACAO[chave]
        for chave in CHAVES_MANIFESTACAO
        if manifestacoes.get(chave)
    ]
    return "; ".join(marcadas) if marcadas else "Nenhuma opcao marcada"


def _determinar_checklist_manifestacao(empresa: dict[str, Any]) -> dict[str, bool]:
    """Determina quais itens do checklist de manifestacao marcar."""
    manifestacoes = empresa.get("manifestacoes")
    if isinstance(manifestacoes, dict):
        normalizadas = normalizar_manifestacoes(manifestacoes, str(empresa.get("manifestacao", "")))
    else:
        normalizadas = inferir_manifestacoes_a_partir_texto(str(empresa.get("manifestacao", "")))

    return {
        "regularizou": normalizadas["regularizou_integralmente"],
        "contestou": normalizadas["apresentou_contestacao"],
        "prorrogacao": normalizadas["solicitou_prorrogacao"],
        "nao_manifestou": normalizadas["nao_apresentou_manifestacao"],
    }


def _gerar_html_individual(empresa: dict[str, Any], auditor: dict[str, str]) -> str:
    """Gera HTML do relatorio individual por CNPJ."""
    checks = _determinar_checklist_manifestacao(empresa)

    def _check_class(ativo: bool) -> str:
        return 'check marcado' if ativo else 'check'

    def _check_x(ativo: bool) -> str:
        return 'X' if ativo else '&nbsp;'

    return f"""<!DOCTYPE html>
<html lang="pt-BR">
<head><meta charset="UTF-8"><style>{CSS_BASE}</style></head>
<body>

{_html_cabecalho()}

<div class="titulo-relatorio">Relatorio Fiscal Conclusivo</div>

<div class="dados-contribuinte">
    <table>
        <tr><td class="label">Contribuinte:</td><td><strong>{empresa.get('contribuinte', '')}</strong></td></tr>
        <tr><td class="label">CNPJ:</td><td>{empresa.get('cnpj', '')}</td></tr>
        <tr><td class="label">Inscricao Estadual:</td><td>{empresa.get('ie', '')}</td></tr>
        <tr><td class="label">DSF:</td><td>{empresa.get('dsf', '')}</td></tr>
        <tr><td class="label">Notificacao DET n.:</td><td>{empresa.get('notificacao_det', '')}</td></tr>
    </table>
</div>

<div class="secao">
    <div class="secao-titulo">1. Identificacao da Acao Fiscal</div>
    <div class="secao-corpo">
        <p>Trata-se de acao fiscal preliminar decorrente de Designacao de Servico Fiscal (DSF),
        no ambito do trabalho de Acervo do sistema Fisconforme, cujo objetivo consistiu na analise
        e cobranca de regularizacao de pendencias fiscais classificadas como prioritarias no painel
        <strong>"Fisconforme Nao Atendido"</strong>.</p>
        <p>As pendencias decorrem de inconsistencias detectadas por malhas fiscais automatizadas
        relativas ao cumprimento de obrigacoes tributarias acessorias.</p>
    </div>
</div>

<div class="secao">
    <div class="secao-titulo">2. Procedimentos Realizados</div>
    <div class="secao-corpo">
        <p>Foram adotadas as seguintes providencias:</p>
        <ul class="lista-procedimentos">
            <li>Consulta ao Sismonitora para identificacao das pendencias e status;</li>
            <li>Verificacao cadastral e fiscal do contribuinte no sistema Visao 360;</li>
            <li>Conferencia do status das inconsistencias no sistema Fisconforme;</li>
            <li>Verificacao da existencia de monitoramentos fiscais relacionados;</li>
            <li>Emissao de notificacao formal via DET, fixando prazo para regularizacao.</li>
        </ul>
    </div>
</div>

<div class="secao">
    <div class="secao-titulo">3. Manifestacao do Contribuinte</div>
    <div class="secao-corpo">
        <p>Apos a notificacao, verificou-se que o contribuinte:</p>
        <div class="manifestacao-box">
            <div class="manifestacao-item"><span class="{_check_class(checks['regularizou'])}">{_check_x(checks['regularizou'])}</span> Regularizou integralmente as pendencias</div>
            <div class="manifestacao-item"><span class="{_check_class(checks['contestou'])}">{_check_x(checks['contestou'])}</span> Apresentou contestacao</div>
            <div class="manifestacao-item"><span class="{_check_class(checks['prorrogacao'])}">{_check_x(checks['prorrogacao'])}</span> Solicitou prorrogacao de prazo</div>
            <div class="manifestacao-item"><span class="{_check_class(checks['nao_manifestou'])}">{_check_x(checks['nao_manifestou'])}</span> Nao apresentou manifestacao</div>
        </div>
        <p>{empresa.get('contatos_realizados', '')}</p>
    </div>
</div>

<div class="secao">
    <div class="secao-titulo">4. Analise Fiscal</div>
    <div class="secao-corpo">
        <p>A analise observou os principios de uniformidade procedimental e vedacao a contestacao
        parcial das pendencias. Foram examinadas as justificativas apresentadas e confrontadas com
        os dados fiscais disponiveis, tendo sido adotadas as seguintes decisoes:</p>
        <div class="decisao-box">{empresa.get('decisao_fiscal', '')}</div>
    </div>
</div>

<div class="secao">
    <div class="secao-titulo">5. Situacao Final das Pendencias</div>
    <div class="secao-corpo">
        <p>Apos a conclusao da acao fiscal, as pendencias apresentaram os seguintes desfechos:</p>
        <div class="desfecho-box">{empresa.get('desfecho', '')}</div>
    </div>
</div>

<div class="secao">
    <div class="secao-titulo">6. Conclusao</div>
    <div class="secao-corpo">
        <p>Diante do exposto, considera-se concluida a presente acao fiscal preliminar, permanecendo
        as pendencias nao regularizadas sujeitas a adocao das medidas fiscais cabiveis, nos termos
        da legislacao vigente.</p>
    </div>
</div>

{_html_assinatura(auditor)}

</body>
</html>"""


def _gerar_html_geral(empresas: list[dict[str, Any]], auditor: dict[str, str]) -> str:
    """Gera HTML do relatorio geral consolidado - formato expandido (3-4 paginas)."""
    total_regularizados = sum(
        1
        for empresa in empresas
        if normalizar_manifestacoes(
            empresa.get("manifestacoes"),
            str(empresa.get("manifestacao", "")),
        ).get("regularizou_integralmente")
    )
    linhas_tabela = ""
    for i, emp in enumerate(empresas, 1):
        linhas_tabela += f"""
        <tr>
            <td style="text-align:center; font-weight:700;">{i}</td>
            <td>{emp.get('contribuinte', '')}</td>
            <td>{emp.get('cnpj', '')}</td>
            <td>{emp.get('notificacao_det', '')}</td>
            <td>{emp.get('manifestacao', '')}</td>
        </tr>"""

    # Blocos de analise individual por contribuinte
    blocos_individuais = ""
    for i, emp in enumerate(empresas, 1):
        blocos_individuais += f"""
<div class="secao" style="margin-top: 8px; page-break-inside: avoid;">
    <div style="font-weight:700; color:#154360; margin-bottom:3px; font-size:10pt;">
        <span class="numero-ordem">{i}</span> {emp.get('contribuinte', '')}
    </div>
    <div class="resumo-individual" style="font-size:9pt;">
        <p><span class="campo-label" style="font-size:9pt;">CNPJ:</span> {emp.get('cnpj', '')}</p>
        <p><span class="campo-label" style="font-size:9pt;">IE:</span> {emp.get('ie', '')}</p>
        <p><span class="campo-label" style="font-size:9pt;">DET n.:</span> {emp.get('notificacao_det', '')}</p>
        <p><span class="campo-label" style="font-size:9pt;">Manifestacao:</span> {emp.get('manifestacao', '')}</p>
    </div>
    <div style="font-size:9pt; margin-top:4px;">
        <div style="font-weight:700; color:#154360; margin-bottom:2px;">Contatos e Manifestacao:</div>
        <p style="margin:2px 0; font-size:8.5pt;">{emp.get('contatos_realizados', '')}</p>
    </div>
    <div style="font-size:9pt; margin-top:4px;">
        <div style="font-weight:700; color:#154360; margin-bottom:2px;">Decisao Fiscal:</div>
        <div class="decisao-box" style="padding:5px 8px; font-size:9pt; font-weight:400;">{emp.get('decisao_fiscal', '')}</div>
    </div>
    <div style="font-size:9pt; margin-top:4px;">
        <div style="font-weight:700; color:#154360; margin-bottom:2px;">Situacao Final:</div>
        <div class="desfecho-box" style="padding:5px 8px; font-size:9pt;">{emp.get('desfecho', '')}</div>
    </div>
</div>"""

    dsf = empresas[0].get("dsf", "") if empresas else ""

    return f"""<!DOCTYPE html>
<html lang="pt-BR">
<head><meta charset="UTF-8"><style>{CSS_BASE}
    .geral-corpo {{ font-size: 10pt; line-height: 1.4; }}
    .geral-corpo .secao {{ margin: 8px 0 4px 0; }}
    .geral-corpo .secao-titulo {{ padding: 4px 10px; font-size: 10.5pt; margin-bottom: 4px; }}
    .geral-corpo .secao-corpo {{ margin: 3px 0 5px 0; }}
    .geral-corpo .secao-corpo p {{ margin: 3px 0; }}
    .geral-corpo .lista-procedimentos li {{ margin-bottom: 2px; font-size: 10pt; }}
    .geral-corpo .tabela-resumo {{ font-size: 9pt; }}
    .geral-corpo .tabela-resumo th {{ padding: 5px 6px; }}
    .geral-corpo .tabela-resumo td {{ padding: 4px 6px; }}
    .geral-corpo .assinatura {{ margin-top: 20px; }}
    .geral-corpo .assinatura .espaco-assinatura {{ height: 70px; }}
    .numero-ordem {{ display: inline-block; background-color: #154360; color: #fff; width: 24px; height: 24px; border-radius: 50%; text-align: center; line-height: 24px; font-size: 10pt; font-weight: 700; margin-right: 6px; vertical-align: middle; }}
    .resumo-individual {{ border: 1px solid #ddd; border-radius: 4px; padding: 8px 10px; margin: 4px 0; background-color: #fafafa; }}
    .resumo-individual p {{ margin: 2px 0; }}
    .resumo-individual .campo-label {{ font-weight: 700; color: #154360; }}
</style></head>
<body class="geral-corpo">

{_html_cabecalho()}

<div class="titulo-relatorio" style="padding:6px 12px; font-size:12pt;">Relatorio Geral Consolidado</div>

<div class="subtitulo-geral" style="font-size:10pt; margin-bottom:8px;">
    Fisconforme Nao Cumprido — DSF n. {dsf}<br>
    Total de contribuintes analisados: <strong>{len(empresas)}</strong>
</div>

<div class="secao">
    <div class="secao-titulo">1. Objeto</div>
    <div class="secao-corpo">
        <p>O presente relatorio consolida os resultados da acao fiscal preliminar realizada no ambito
        da Designacao de Servico Fiscal (DSF) n. {dsf}, referente ao trabalho de Acervo do sistema
        Fisconforme, com foco nas pendencias classificadas como prioritarias no painel
        <strong>"Fisconforme Nao Atendido"</strong>.</p>
    </div>
</div>

<div class="secao">
    <div class="secao-titulo">2. Quadro Resumo dos Contribuintes</div>
    <div class="secao-corpo">
        <table class="tabela-resumo">
            <thead>
                <tr>
                    <th style="width:24px; text-align:center;">N.</th>
                    <th>Contribuinte</th>
                    <th style="width:115px;">CNPJ</th>
                    <th style="width:70px;">DET n.</th>
                    <th style="width:135px;">Manifestacao</th>
                </tr>
            </thead>
            <tbody>{linhas_tabela}</tbody>
        </table>
    </div>
</div>

<div class="secao">
    <div class="secao-titulo">3. Procedimentos Comuns</div>
    <div class="secao-corpo">
        <p>Para todos os contribuintes listados, foram adotadas as seguintes providencias:</p>
        <ul class="lista-procedimentos">
            <li>Consulta ao Sismonitora para identificacao das pendencias e status;</li>
            <li>Verificacao cadastral e fiscal do contribuinte no sistema Visao 360;</li>
            <li>Conferencia do status das inconsistencias no sistema Fisconforme;</li>
            <li>Verificacao da existencia de monitoramentos fiscais relacionados;</li>
            <li>Emissao de notificacao formal via DET, fixando prazo para regularizacao;</li>
            <li>Tentativa de contato por e-mail, WhatsApp e telefone.</li>
        </ul>
    </div>
</div>

<div class="secao">
    <div class="secao-titulo">4. Analise Individual por Contribuinte</div>
</div>

{blocos_individuais}

<div style="page-break-before: always;">
        <div class="secao">
            <div class="secao-titulo">5. Conclusao Geral</div>
        <div class="secao-corpo">
            <p>Diante do exposto, dos <strong>{len(empresas)}</strong> contribuintes notificados no
            ambito da DSF n. {dsf}, verificou-se o seguinte panorama:</p>
            <ul class="lista-procedimentos">
                <li>{total_regularizados} promoveram a regularização integral das pendências
                apontadas pelo sistema Fisconforme, tendo suas situações sido regularizadas;</li>
                <li>01 (um) contribuinte apresentou contestacao formal; contudo, a resolucao
                da pendencia depende do registro de evento de cancelamento de NF-e, o que
                permanece pendente ate o encerramento desta analise;</li>
                <li>01 (um) contribuinte solicitou prorrogacao de prazo para regularizacao,
                sem que tenha sido verificado o efetivo saneamento das pendencias dentro do novo prazo concedido;</li>
                <li>Os demais contribuintes mantiveram-se inertes diante dos contatos
                realizados e das notificacoes formais enviadas via DET, nao apresentando qualquer
                manifestacao ou providencia de regularizacao.</li>
            </ul>
            <p>Recomenda-se o encaminhamento para as providencias fiscais cabiveis, nos termos da
            legislacao vigente, em relacao aos contribuintes que permanecem com pendencias ativas.</p>
        </div>
    </div>
    {_html_assinatura(auditor)}
</div>

</body>
</html>"""


def _mesclar_pdfs(pdfs: list[str], output_path: str) -> None:
    """Mescla uma lista de PDFs em um unico arquivo."""
    PdfWriter = _importar_pdf_writer()
    merger = PdfWriter()
    for pdf in pdfs:
        if os.path.exists(pdf):
            merger.append(pdf)
        else:
            logger.warning("PDF nao encontrado para mesclagem: %s", pdf)
    with open(output_path, "wb") as fout:
        merger.write(fout)


def _extrair_texto_simples_html(html_content: str) -> list[str]:
    """Converte o HTML gerado em linhas simples para o fallback em reportlab."""
    html_sem_estilos = re.sub(r"(?is)<style.*?>.*?</style>", "", html_content)
    html_com_quebras = re.sub(r"(?i)</(p|div|li|tr|h1|h2|h3|table|ul)>", "\n", html_sem_estilos)
    texto = re.sub(r"(?s)<[^>]+>", "", html_com_quebras)
    texto = (
        texto.replace("&nbsp;", " ")
        .replace("&mdash;", "-")
        .replace("&amp;", "&")
    )
    return [linha.strip() for linha in texto.splitlines() if linha.strip()]


def _renderizar_pdf_com_reportlab(html_content: str, output_path: str) -> None:
    """Gera um PDF textual com reportlab quando o WeasyPrint nao estiver disponivel."""
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    largura, altura = A4
    pdf = canvas.Canvas(output_path, pagesize=A4)
    y = altura - 50

    for linha in _extrair_texto_simples_html(html_content):
        if y < 50:
            pdf.showPage()
            y = altura - 50
        pdf.drawString(40, y, linha[:110])
        y -= 14

    pdf.save()


def _renderizar_pdf_html_ou_fallback(html_content: str, output_path: str) -> None:
    """Renderiza o PDF via WeasyPrint e usa fallback puro Python quando necessario."""
    try:
        HTML = _importar_weasyprint_html()
        HTML(string=html_content).write_pdf(output_path)
        return
    except Exception as erro_weasy:  # noqa: BLE001
        logger.warning("WeasyPrint indisponivel, usando fallback reportlab: %s", erro_weasy)

    if importlib.util.find_spec("reportlab") is None:
        raise RuntimeError("Nenhum renderizador PDF disponivel: instale weasyprint ou reportlab.")

    _renderizar_pdf_com_reportlab(html_content, output_path)


def _normalizar_lista_pdfs_notificacao(pdfs_notificacao: Optional[Any]) -> list[str]:
    """Normaliza a colecao de PDFs de notificacao recebida pela API."""
    if not pdfs_notificacao:
        return []
    if isinstance(pdfs_notificacao, str):
        return [pdfs_notificacao]
    if isinstance(pdfs_notificacao, list):
        return [str(pdf) for pdf in pdfs_notificacao]
    return []


# ============================================================
# FUNCOES DE PERSISTENCIA
# ============================================================

def carregar_json(caminho: Path, padrao: Any = None) -> Any:
    """Carrega JSON de disco com fallback para valor padrao."""
    if padrao is None:
        padrao = {}
    if not caminho.exists():
        return padrao
    with caminho.open("r", encoding="utf-8") as arquivo:
        return json.load(arquivo)


def salvar_json(caminho: Path, payload: Any) -> None:
    """Salva JSON em disco com indentacao padrao."""
    caminho.parent.mkdir(parents=True, exist_ok=True)
    with caminho.open("w", encoding="utf-8") as arquivo:
        json.dump(payload, arquivo, ensure_ascii=False, indent=2)


def obter_caminho_auditor(base_dir: Path) -> Path:
    """Resolve caminho do arquivo de dados do auditor."""
    return base_dir.parent / "_config" / "auditor.json"


def obter_caminho_dsf(base_dir: Path) -> Path:
    """Resolve caminho do arquivo de DSFs com CNPJs vinculados."""
    return base_dir.parent / "_config" / "dsf.json"


def obter_caminho_relatorio_cnpj(base_dir: Path, cnpj: str) -> Path:
    """Resolve caminho do arquivo de dados do relatorio por CNPJ."""
    return base_dir / cnpj / "relatorio" / "dados.json"


def obter_diretorio_relatorio_cnpj(base_dir: Path, cnpj: str) -> Path:
    """Resolve diretorio de relatorio do CNPJ."""
    diretorio = base_dir / cnpj / "relatorio"
    diretorio.mkdir(parents=True, exist_ok=True)
    return diretorio


# ============================================================
# DSF - Gerenciamento de CNPJs vinculados
# ============================================================

def carregar_dsfs(base_dir: Path) -> dict[str, Any]:
    """Carrega todas as DSFs com seus CNPJs vinculados.

    Formato:
    {
        "20263710400226": {
            "numero": "20263710400226",
            "descricao": "Acervo Fisconforme",
            "cnpjs": ["23722199000167", "46675389000176", ...]
        }
    }
    """
    return carregar_json(obter_caminho_dsf(base_dir), {})


def salvar_dsfs(base_dir: Path, dsfs: dict[str, Any]) -> None:
    """Salva todas as DSFs."""
    salvar_json(obter_caminho_dsf(base_dir), dsfs)


def obter_dsf(base_dir: Path, numero_dsf: str) -> dict[str, Any] | None:
    """Obtem uma DSF especifica pelo numero."""
    dsfs = carregar_dsfs(base_dir)
    return dsfs.get(numero_dsf)


def salvar_dsf(base_dir: Path, numero_dsf: str, descricao: str, cnpjs: list[str]) -> dict[str, Any]:
    """Salva ou atualiza uma DSF com seus CNPJs vinculados."""
    dsfs = carregar_dsfs(base_dir)
    dsfs[numero_dsf] = {
        "numero": numero_dsf,
        "descricao": descricao,
        "cnpjs": cnpjs,
    }
    salvar_dsfs(base_dir, dsfs)
    return dsfs[numero_dsf]


def listar_dsfs(base_dir: Path) -> list[dict[str, Any]]:
    """Lista todas as DSFs cadastradas."""
    dsfs = carregar_dsfs(base_dir)
    return list(dsfs.values())


def carregar_dsfs_efetivas(base_dir: Path) -> dict[str, Any]:
    """Mescla DSFs persistidas com as inferidas dos relatorios por CNPJ."""
    dsfs_persistidas = carregar_dsfs(base_dir)
    dsfs_mescladas = {
        numero: {
            "numero": dados.get("numero", numero),
            "descricao": dados.get("descricao", ""),
            "cnpjs": sorted({re.sub(r"\D", "", cnpj) for cnpj in dados.get("cnpjs", []) if len(re.sub(r"\D", "", cnpj)) == 14}),
            "origem": "persistida",
        }
        for numero, dados in dsfs_persistidas.items()
    }

    for item in listar_cnpjs_com_relatorio(base_dir):
        numero_dsf = item.get("dsf", "")
        cnpj = item.get("cnpj", "")
        if not numero_dsf or not cnpj:
            continue
        if numero_dsf not in dsfs_mescladas:
            dsfs_mescladas[numero_dsf] = {
                "numero": numero_dsf,
                "descricao": "Acervo Fisconforme",
                "cnpjs": [],
                "origem": "inferida",
            }
        if cnpj not in dsfs_mescladas[numero_dsf]["cnpjs"]:
            dsfs_mescladas[numero_dsf]["cnpjs"].append(cnpj)

    for dados in dsfs_mescladas.values():
        dados["cnpjs"] = sorted(dados["cnpjs"])
        if dados["origem"] == "persistida" and dados["cnpjs"]:
            dados["origem"] = "mesclada"

    return dsfs_mescladas


def _listar_dets_locais(base_dir: Path, cnpj: str) -> list[dict[str, Any]]:
    """Lista PDFs de notificacao presentes localmente para um CNPJ."""
    diretorio = obter_diretorio_relatorio_cnpj(base_dir, cnpj)
    dados_brutos = carregar_json(obter_caminho_relatorio_cnpj(base_dir, cnpj), {})
    numero_notificacao = re.sub(r"\D+", "", str(dados_brutos.get("notificacao_det", "")))
    dets: list[dict[str, Any]] = []

    for arquivo in sorted(diretorio.iterdir()):
        if not arquivo.is_file():
            continue
        if arquivo.suffix.lower() != ".pdf":
            continue

        nome_normalizado = _normalizar_texto_busca(arquivo.name)
        digitos_arquivo = re.sub(r"\D+", "", arquivo.stem)
        eh_relatorio_final = "relatorio_final" in nome_normalizado or "relatorio_geral" in nome_normalizado
        menciona_det = "det" in nome_normalizado
        menciona_notificacao = "notific" in nome_normalizado
        corresponde_numero_notificacao = bool(numero_notificacao and numero_notificacao in digitos_arquivo)

        # Ignorar os PDFs finais ja gerados para manter somente os anexos de notificacao.
        if eh_relatorio_final:
            continue

        # Aceitar nomes antigos contendo "det", arquivos de notificacao e anexos
        # identificados pelo numero salvo no relatorio.
        if not (menciona_det or menciona_notificacao or corresponde_numero_notificacao):
            continue

        dets.append(
            {
                "nome": arquivo.name,
                "caminho": str(arquivo),
                "tamanho_bytes": arquivo.stat().st_size,
            }
        )

    return dets


def _normalizar_arquivos_notificacao_incluidos(
    arquivos_notificacao_incluidos: Any,
    dets_locais: list[dict[str, Any]],
) -> list[str]:
    """Normaliza a selecao de PDFs usando apenas nomes locais validos."""
    nomes_disponiveis = [det["nome"] for det in dets_locais]
    if not nomes_disponiveis:
        return []
    if arquivos_notificacao_incluidos is None or not isinstance(arquivos_notificacao_incluidos, list):
        return nomes_disponiveis
    return [nome for nome in arquivos_notificacao_incluidos if nome in nomes_disponiveis]


def listar_dets_disponiveis(
    base_dir: Path,
    cnpj: str,
    arquivos_notificacao_incluidos: Any = _SENTINELA_SELECAO_AUSENTE,
) -> list[dict[str, Any]]:
    """Lista PDFs locais com o estado de selecao persistido."""
    dets_locais = _listar_dets_locais(base_dir, cnpj)
    if arquivos_notificacao_incluidos is _SENTINELA_SELECAO_AUSENTE:
        dados_brutos = carregar_json(obter_caminho_relatorio_cnpj(base_dir, cnpj), {})
        arquivos_notificacao_incluidos = (
            dados_brutos["arquivos_notificacao_incluidos"]
            if "arquivos_notificacao_incluidos" in dados_brutos
            else None
        )

    nomes_selecionados = set(
        _normalizar_arquivos_notificacao_incluidos(
            arquivos_notificacao_incluidos,
            dets_locais,
        )
    )
    return [{**det, "selecionado": det["nome"] in nomes_selecionados} for det in dets_locais]


def resolver_caminho_det(
    base_dir: Path,
    cnpj: str,
    caminho_informado: Optional[str],
) -> tuple[Optional[str], Optional[str]]:
    """Resolve o caminho legado `pdf_det` ou registra o motivo da indisponibilidade."""
    if caminho_informado:
        caminho_det = Path(caminho_informado)
        if caminho_det.exists() and caminho_det.is_file():
            return str(caminho_det), None
        aviso = f"DET informado nao encontrado no caminho configurado: {caminho_informado}"
    else:
        aviso = None

    dets_locais = _listar_dets_locais(base_dir, cnpj)
    if dets_locais:
        return dets_locais[0]["caminho"], None
    return None, aviso


def carregar_dados_relatorio(base_dir: Path, cnpj: str) -> dict[str, Any]:
    """Carrega e normaliza os dados do relatorio do CNPJ."""
    dados_brutos = carregar_json(obter_caminho_relatorio_cnpj(base_dir, cnpj), {})
    manifestacoes = normalizar_manifestacoes(
        dados_brutos.get("manifestacoes"),
        str(dados_brutos.get("manifestacao", "")),
    )

    dets_disponiveis = _listar_dets_locais(base_dir, cnpj)
    possui_selecao_persistida = "arquivos_notificacao_incluidos" in dados_brutos
    arquivos_notificacao_incluidos = _normalizar_arquivos_notificacao_incluidos(
        dados_brutos["arquivos_notificacao_incluidos"] if possui_selecao_persistida else None,
        dets_disponiveis,
    )
    caminhos_notificacao_incluidos = [
        det["caminho"]
        for det in dets_disponiveis
        if det["nome"] in arquivos_notificacao_incluidos
    ]

    pdf_det_legado, aviso_det = resolver_caminho_det(base_dir, cnpj, dados_brutos.get("pdf_det"))
    if not possui_selecao_persistida and not caminhos_notificacao_incluidos and pdf_det_legado:
        caminhos_notificacao_incluidos = [pdf_det_legado]
        nome_legado = Path(pdf_det_legado).name
        if nome_legado in [det["nome"] for det in dets_disponiveis]:
            arquivos_notificacao_incluidos = [nome_legado]

    dados_normalizados = {
        "cnpj": dados_brutos.get("cnpj") or cnpj,
        "contribuinte": dados_brutos.get("contribuinte", ""),
        "ie": dados_brutos.get("ie", ""),
        "dsf": dados_brutos.get("dsf", ""),
        "notificacao_det": dados_brutos.get("notificacao_det", ""),
        "manifestacoes": manifestacoes,
        "manifestacao": resumir_manifestacoes(manifestacoes),
        "contatos_realizados": dados_brutos.get("contatos_realizados", ""),
        "decisao_fiscal": dados_brutos.get("decisao_fiscal", ""),
        "desfecho": dados_brutos.get("desfecho", ""),
        "arquivos_notificacao_incluidos": arquivos_notificacao_incluidos,
        "caminhos_notificacao_incluidos": caminhos_notificacao_incluidos,
        "pdf_det": caminhos_notificacao_incluidos[0] if caminhos_notificacao_incluidos else None,
        "tem_det": bool(caminhos_notificacao_incluidos),
    }
    if aviso_det:
        dados_normalizados["aviso_det"] = aviso_det
    return dados_normalizados


def listar_cnpjs_com_relatorio(base_dir: Path) -> list[dict[str, Any]]:
    """Lista CNPJs com relatorio preenchido e diagnostico do DET."""
    if not base_dir.exists():
        return []

    cnpjs_com_relatorio: list[dict[str, Any]] = []
    for item in sorted(base_dir.iterdir()):
        if not item.is_dir() or item.name.startswith("_"):
            continue
        dados_path = item / "relatorio" / "dados.json"
        if not dados_path.exists():
            continue
        dados = carregar_dados_relatorio(base_dir, item.name)
        if not dados.get("contribuinte"):
            continue
        cnpjs_com_relatorio.append(
            {
                "cnpj": item.name,
                "contribuinte": dados.get("contribuinte", ""),
                "dsf": dados.get("dsf", ""),
                "notificacao_det": dados.get("notificacao_det", ""),
                "tem_det": dados.get("tem_det", False),
                "manifestacao": dados.get("manifestacao", ""),
                "aviso_det": dados.get("aviso_det"),
            }
        )
    return cnpjs_com_relatorio


def diagnosticar_prontidao_relatorios(base_dir: Path) -> dict[str, Any]:
    """Resume dependencias e artefatos necessarios para gerar relatorios."""
    dependencias = [
        _diagnosticar_dependencia_weasyprint(),
        _diagnosticar_dependencia_reportlab(),
        _diagnosticar_dependencia_pypdf(),
        _diagnosticar_dependencia_python_docx(),
    ]
    renderizador_disponivel = any(
        item["nome"] in {"weasyprint", "reportlab"} and item["instalado"]
        for item in dependencias
    )
    pypdf_disponivel = any(item["nome"] == "pypdf" and item["instalado"] for item in dependencias)
    docx_disponivel = any(item["nome"] == "python-docx" and item["instalado"] for item in dependencias)
    return {
        "pronto_pdf": renderizador_disponivel and pypdf_disponivel and docx_disponivel,
        "dependencias": dependencias,
        "dependencias_faltantes": [item["nome"] for item in dependencias if not item["instalado"]],
        "modelos_docx": diagnosticar_modelos_docx(),
        "total_cnpjs_com_relatorio": len(listar_cnpjs_com_relatorio(base_dir)),
        "total_dsfs": len(carregar_dsfs_efetivas(base_dir)),
        "cnpjs_com_relatorio": listar_cnpjs_com_relatorio(base_dir),
        "dsfs": list(carregar_dsfs_efetivas(base_dir).values()),
    }


# ============================================================
# GERACAO DE PDFs
# ============================================================

def gerar_pdf_individual(
    empresa: dict[str, Any],
    auditor: dict[str, str],
    output_path: str,
    pdfs_notificacao: Optional[Any] = None,
) -> str:
    """Gera PDF do relatorio individual e opcionalmente mescla com DET."""
    html_content = _gerar_html_individual(empresa, auditor)

    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        temp_pdf = tmp.name

    try:
        _renderizar_pdf_html_ou_fallback(html_content, temp_pdf)
        pdfs_para_unir = [temp_pdf]
        for pdf in _normalizar_lista_pdfs_notificacao(pdfs_notificacao):
            if os.path.exists(pdf):
                pdfs_para_unir.append(pdf)
        if len(pdfs_para_unir) > 1:
            _mesclar_pdfs(pdfs_para_unir, output_path)
        else:
            shutil.copy2(temp_pdf, output_path)
    finally:
        if os.path.exists(temp_pdf):
            os.remove(temp_pdf)

    return output_path


def gerar_pdf_geral(
    empresas: list[dict[str, Any]],
    auditor: dict[str, str],
    output_path: str,
    incluir_dets: bool = True,
) -> str:
    """Gera PDF do relatorio geral consolidado e opcionalmente mescla com DETs."""
    html_content = _gerar_html_geral(empresas, auditor)

    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        temp_pdf = tmp.name

    try:
        _renderizar_pdf_html_ou_fallback(html_content, temp_pdf)

        pdfs_para_unir = [temp_pdf]
        if incluir_dets:
            for emp in empresas:
                caminhos = emp.get("caminhos_notificacao_incluidos") or []
                if not caminhos and emp.get("pdf_det"):
                    caminhos = [emp["pdf_det"]]
                for det in caminhos:
                    if det and os.path.exists(det):
                        pdfs_para_unir.append(det)

        if len(pdfs_para_unir) > 1:
            _mesclar_pdfs(pdfs_para_unir, output_path)
        else:
            shutil.copy2(temp_pdf, output_path)
    finally:
        if os.path.exists(temp_pdf):
            os.remove(temp_pdf)

    return output_path
